"""
Web interface routes for FlowSlide
"""

import asyncio
import json
import logging
import os
import re
import tempfile
import time
import urllib.parse
import uuid
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional
import hashlib

from bs4 import BeautifulSoup
from fastapi import APIRouter, Depends, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, StreamingResponse, JSONResponse
from fastapi.templating import Jinja2Templates
from pydantic import BaseModel
from sqlalchemy.orm import Session

from ..ai import AIMessage, MessageRole, get_ai_provider
from ..api.models import FileOutlineGenerationRequest, PPTGenerationRequest
from ..auth.middleware import get_current_user_optional, get_current_user_required
from ..core.simple_config import ai_config
from ..database.database import get_db
from ..database.models import User
from ..services.pdf_to_pptx_converter import get_pdf_to_pptx_converter
from ..services.pyppeteer_pdf_converter import get_pdf_converter
from ..utils.thread_pool import run_blocking_io

# Import shared service instances to ensure data consistency
from ..services.service_instances import ppt_service

# Configure logger for this module
logger = logging.getLogger(__name__)


# Simple filter to allow only records at or above a given level
class _LevelFilter(logging.Filter):
    def __init__(self, level=logging.WARNING):
        super().__init__()
        self._level = level

    def filter(self, record):
        try:
            return record.levelno >= self._level
        except Exception:
            return True


# Helper function for aspect ratio settings
def get_aspect_ratio_settings() -> dict:
    """Get aspect ratio related settings from config (defaults to 16:9)."""
    try:
        from ..services.config_service import get_config_service

        cfg = get_config_service().get_all_config()
        aspect = (cfg.get("aspect_ratio") or "16:9").strip()
    except Exception:
        aspect = "16:9"

    if aspect == "4:3":
        return {"ratio_css": "4/3", "width": 1280, "height": 960, "vw_height": "75vw"}
    # default 16:9
    return {"ratio_css": "16/9", "width": 1280, "height": 720, "vw_height": "56.25vw"}


router = APIRouter()
@router.get("/api/research/reports")
async def list_research_reports(user: User = Depends(get_current_user_required)):
    """列出本地 research_reports 下的已生成研究报告（Markdown/Txt）。"""
    from pathlib import Path
    reports_dir = Path("research_reports")
    items = []
    try:
        if reports_dir.exists():
            for p in sorted(reports_dir.glob("*.md"), reverse=True):
                try:
                    head = p.read_text(encoding="utf-8", errors="ignore").splitlines()[:8]
                    items.append(
                        {
                            "filename": p.name,
                            "size": p.stat().st_size,
                            "modified": int(p.stat().st_mtime),
                            "preview": "\n".join(head),
                        }
                    )
                except Exception:
                    # 单个文件读取失败忽略
                    continue
            # 也允许 .txt
            for p in sorted(reports_dir.glob("*.txt"), reverse=True):
                if any(x["filename"] == p.name for x in items):
                    continue
                try:
                    head = p.read_text(encoding="utf-8", errors="ignore").splitlines()[:8]
                    items.append(
                        {
                            "filename": p.name,
                            "size": p.stat().st_size,
                            "modified": int(p.stat().st_mtime),
                            "preview": "\n".join(head),
                        }
                    )
                except Exception:
                    continue
    except Exception as e:
        return {"success": False, "error": str(e), "reports": []}
    return {"success": True, "reports": items}

@router.get("/api/research/reports/{filename}")
async def get_research_report(filename: str, user: User = Depends(get_current_user_required)):
    """获取指定研究报告全文内容(安全限制)。"""
    try:
        safe = filename.replace("..", "").strip("/\\")
        if not safe or safe != filename:
            raise HTTPException(status_code=400, detail="非法文件名")
        from pathlib import Path
        reports_dir = Path("research_reports")
        target = reports_dir / safe
        if not (target.exists() and target.is_file()):
            raise HTTPException(status_code=404, detail="文件不存在")
        if target.suffix.lower() not in {".md", ".txt"}:
            raise HTTPException(status_code=400, detail="不支持的文件类型")
        content = target.read_text(encoding="utf-8", errors="ignore")
        max_len = 300_000
        truncated = False
        if len(content) > max_len:
            content = content[:max_len]
            truncated = True
        return {
            "success": True,
            "filename": safe,
            "size": target.stat().st_size,
            "modified": int(target.stat().st_mtime),
            "truncated": truncated,
            "content": content,
        }
    except HTTPException:
        raise
    except Exception as e:
        return {"success": False, "error": str(e)}

def _sanitize_report_filename(name: str) -> str:
    base = name.replace('..','').replace('\r','').replace('\n','').strip().strip('/\\')
    if not base:
        base = f"report_{int(time.time())}.md"
    # 只允许字母数字下划线中划线点和中文
    import re
    base = re.sub(r"[^A-Za-z0-9_.\-\u4e00-\u9fa5]", "_", base)
    # 强制后缀
    if not base.lower().endswith(('.md','.txt')):
        base += '.md'
    return base[:120]

@router.get("/api/research/reports/{filename}/chunk")
async def get_research_report_chunk(
    filename: str,
    offset: int = 0,
    limit: int = 20000,
    user: User = Depends(get_current_user_required)
):
    """分页获取报告内容，便于前端懒加载。"""
    try:
        safe = filename.replace("..", "").strip("/\\")
        from pathlib import Path
        reports_dir = Path("research_reports")
        target = reports_dir / safe
        if not (target.exists() and target.is_file()):
            raise HTTPException(status_code=404, detail="文件不存在")
        if target.suffix.lower() not in {".md", ".txt"}:
            raise HTTPException(status_code=400, detail="不支持的文件类型")
        data = target.read_text(encoding="utf-8", errors="ignore")
        size = len(data)
        if offset < 0: offset = 0
        if limit <= 0: limit = 20000
        if limit > 50000: limit = 50000
        chunk = data[offset: offset+limit]
        next_offset = offset + len(chunk)
        has_more = next_offset < size
        return {"success": True, "filename": safe, "offset": offset, "next_offset": next_offset, "has_more": has_more, "size": size, "content": chunk}
    except HTTPException:
        raise
    except Exception as e:
        return {"success": False, "error": str(e)}

@router.post("/api/research/reports/upload")
async def upload_research_report(
    file: UploadFile = File(...),
    user: User = Depends(get_current_user_required)
):
    """单独上传并缓存研究报告到 research_reports 目录，供后续复用。"""
    try:
        fname = file.filename or "uploaded_report.md"
        if not any(fname.lower().endswith(ext) for ext in [".md", ".txt"]):
            raise HTTPException(status_code=400, detail="仅支持 .md/.txt")
        raw = await file.read()
        text = raw.decode("utf-8", errors="ignore")
        stored_name = _sanitize_report_filename(fname)
        from pathlib import Path
        reports_dir = Path("research_reports")
        reports_dir.mkdir(exist_ok=True)
        # 若存在则加时间戳避免覆盖
        candidate = stored_name
        if (reports_dir / candidate).exists():
            stem, ext = os.path.splitext(candidate)
            candidate = f"{stem}_{int(time.time())}{ext}"
        path = reports_dir / candidate
        # 限制文件大小写入（最多 400k 字符）
        truncated = False
        if len(text) > 400_000:
            text = text[:400_000]
            truncated = True
        path.write_text(text, encoding="utf-8")
        preview = "\n".join(text.splitlines()[:8])
        return {"success": True, "filename": candidate, "size": len(text), "truncated": truncated, "preview": preview}
    except HTTPException:
        raise
    except Exception as e:
        return {"success": False, "error": str(e)}


def _build_provider_status(ai_config_obj) -> dict:
    """Return a dict mapping provider->availability to avoid long inline comprehensions."""
    try:
        providers = ai_config_obj.get_available_providers()
        status = {}
        for p in providers:
            try:
                status[p] = ai_config_obj.is_provider_available(p)
            except Exception:
                status[p] = False
        return status
    except Exception:
        return {}
# Templates directory - use absolute path for better reliability


# Templates directory - use absolute path for better reliability

template_dir = os.path.join(os.path.dirname(__file__), "templates")
templates = Jinja2Templates(directory=template_dir)
def _redact_key(value: str, keep: int = 4) -> str:
    try:
        s = str(value or "")
        if not s:
            return s
        if len(s) <= keep:
            return "*" * len(s)
        return s[:keep] + "…" + "*" * max(0, len(s) - keep - 1)
    except Exception:
        return "***"


def _sanitize_text(text: str) -> str:
    """Remove or mask likely API keys and tokens in an error text.

    - Masks patterns like key=XXXX, Authorization: Bearer XXXX, X-Api-Key: XXXX
    - Truncates very long bodies
    """
    try:
        if not text:
            return text
        t = str(text)
        # Basic masks
        import re

        t = re.sub(r"(key=)([^&\s]{8,})", lambda m: m.group(1) + _redact_key(m.group(2)), t)
        t = re.sub(r"(api[_-]?key[\"']?[:=]\s*)([A-Za-z0-9._-]{8,})", lambda m: m.group(1) + _redact_key(m.group(2)), t, flags=re.IGNORECASE)
        t = re.sub(r"(Authorization:\s*Bearer\s+)([A-Za-z0-9._-]{8,})", lambda m: m.group(1) + _redact_key(m.group(2)), t, flags=re.IGNORECASE)
        t = re.sub(r"(X-Api-Key:\s*)([A-Za-z0-9._-]{8,})", lambda m: m.group(1) + _redact_key(m.group(2)), t, flags=re.IGNORECASE)

        # Trim output to avoid dumping entire HTML/JSON
        if len(t) > 800:
            t = t[:800] + "…"
        return t
    except Exception:
        return "(error text hidden)"


def _sanitize_dict(d: dict) -> dict:
    try:
        if not isinstance(d, dict):
            return d
        redacted = {}
        for k, v in d.items():
            key_l = str(k).lower()
            if any(s in key_l for s in ["api_key", "apikey", "x-api-key", "authorization", "token", "secret"]):
                redacted[k] = _redact_key(str(v))
            else:
                redacted[k] = v
        return redacted
    except Exception:
        return {"info": "(payload hidden)"}



# Helper to join base URL and endpoint paths safely
def build_api_url(base_url: str, *parts: str, ensure_v1: bool = False) -> str:
    """Return a safely-joined URL.

    - If ensure_v1 is True, ensures base_url ends with '/v1' before joining.
    - parts are appended without duplicate slashes.
    """
    if not base_url:
        return "/" + "/".join(p.strip("/") for p in parts)

    base = base_url.rstrip("/")
    if ensure_v1 and not base.endswith("/v1"):
        base = base + "/v1"

    suffix = "/".join(p.strip("/") for p in parts if p)
    return f"{base}/{suffix}" if suffix else base


# Add custom filters
def timestamp_to_datetime(timestamp):
    """Convert timestamp to readable datetime string"""
    try:
        if isinstance(timestamp, (int, float)):
            return datetime.fromtimestamp(timestamp).strftime("%Y-%m-%d %H:%M:%S")
        return str(timestamp)
    except (ValueError, OSError):
        return "无效时间"


def strftime_filter(timestamp, format_string="%Y-%m-%d %H:%M"):
    """Jinja2 strftime filter"""
    try:
        if isinstance(timestamp, (int, float)):
            dt = datetime.fromtimestamp(timestamp)
            return dt.strftime(format_string)
        return str(timestamp)
    except (ValueError, OSError):
        return "无效时间"


# Register custom filters
templates.env.filters["timestamp_to_datetime"] = timestamp_to_datetime
templates.env.filters["strftime"] = strftime_filter


# --- RBAC helpers ---
def _user_can_access_project(user: Optional[User], project: Any) -> bool:
    try:
        # If no user, allow access for testing purposes (temporary)
        if user is None:
            return True

        if getattr(user, "is_admin", False):
            return True
        owner_id = getattr(project, "owner_id", None)
        # Strict: only owner can access; legacy projects with owner_id=None are admin-only
        return owner_id is not None and owner_id == getattr(user, "id", None)
    except Exception:
        return False


# --- Admin: Simple user management page ---
@router.get("/admin/users", response_class=HTMLResponse)
async def admin_users_page(request: Request, user: User = Depends(get_current_user_required)):
    """简易管理员用户管理页面（基于现有 JSON API）"""
    # 仅管理员可访问
    if not user.is_admin:
        from fastapi.responses import RedirectResponse

        return RedirectResponse(url="/home", status_code=302)

    return templates.TemplateResponse(
        "admin_users.html",
        {
            "request": request,
            "user": user.to_dict(),
        },
    )


# AI编辑请求数据模型
class AISlideEditRequest(BaseModel):
    slideIndex: int
    slideTitle: str
    slideContent: str
    userRequest: str
    projectInfo: Dict[str, Any]
    slideOutline: Optional[Dict[str, Any]] = None
    chatHistory: Optional[List[Dict[str, str]]] = None
    images: Optional[List[Dict[str, str]]] = None  # 新增：图片信息列表


# AI要点增强请求数据模型
class AIBulletPointEnhanceRequest(BaseModel):
    slideIndex: int
    slideTitle: str
    slideContent: str
    userRequest: str
    projectInfo: Dict[str, Any]
    slideOutline: Optional[Dict[str, Any]] = None
    contextInfo: Optional[Dict[str, Any]] = None  # 包含原始要点、其他要点等上下文信息


# 简单通用补全请求（用于演讲稿润色等轻量场景）
class SimpleCompletionRequest(BaseModel):
    prompt: str
    max_tokens: Optional[int] = None
    temperature: Optional[float] = None
    provider: Optional[str] = None  # 预留：未来可切换不同模型


# 图像重新生成请求数据模型
class AIImageRegenerateRequest(BaseModel):
    slide_index: int
    image_info: Dict[str, Any]
    slide_content: Dict[str, Any]
    project_topic: str
    project_scenario: str
    regeneration_reason: Optional[str] = None


# 一键配图请求数据模型
class AIAutoImageGenerateRequest(BaseModel):
    slide_index: int
    slide_content: Dict[str, Any]
    project_topic: str
    project_scenario: str


# Helper function to extract slides from HTML content
async def _extract_slides_from_html(slides_html: str, existing_slides_data: list) -> list:
    """
    Extract individual slides from combined HTML content and update slides_data
    """
    try:
        # Parse HTML content
        soup = BeautifulSoup(slides_html, "html.parser")

        # Find all slide containers - look for common slide patterns
        slide_containers = []

        # Try different patterns to find slides
        patterns = [
            {"class": re.compile(r"slide")},
            {"class": re.compile(r"page")},
            {"style": re.compile(r"width:\s*1280px.*height:\s*720px", re.IGNORECASE)},
            {"style": re.compile(r"aspect-ratio:\s*16\s*/\s*9", re.IGNORECASE)},
        ]

        for pattern in patterns:
            containers = soup.find_all("div", pattern)
            if containers:
                slide_containers = containers
                break

        # If no specific slide containers found, try to split by common separators
        if not slide_containers:
            # Look for sections or divs that might represent slides
            all_divs = soup.find_all("div")
            # Filter divs that might be slides (have substantial content)
            slide_containers = [
                div
                for div in all_divs
                if div.get_text(strip=True) and len(div.get_text(strip=True)) > 50
            ]

        updated_slides_data = []

        # If we found slide containers, extract them
        if slide_containers:
            for i, container in enumerate(slide_containers):
                # Try to extract title from the slide
                title = f"第{i+1}页"
                title_elements = container.find_all(["h1", "h2", "h3", "h4", "h5", "h6"])
                if title_elements:
                    title = title_elements[0].get_text(strip=True) or title

                # Get the HTML content of this slide
                slide_html = str(container)

                # Create slide data
                slide_data = {
                    "page_number": i + 1,
                    "title": title,
                    "html_content": slide_html,
                    "is_user_edited": True,  # Mark as user edited since it came from editor
                }

                # If we have existing slide data, preserve some fields
                if i < len(existing_slides_data):
                    existing_slide = existing_slides_data[i]
                    # Preserve any additional fields from existing data
                    for key, value in existing_slide.items():
                        if key not in slide_data:
                            slide_data[key] = value

                updated_slides_data.append(slide_data)

        # If we couldn't extract individual slides, treat the entire content as slides
        if not updated_slides_data and existing_slides_data:
            # Fall back to using existing slides structure but mark as edited
            for i, existing_slide in enumerate(existing_slides_data):
                slide_data = existing_slide.copy()
                slide_data["is_user_edited"] = True
                updated_slides_data.append(slide_data)

        # If we still have no slides but have HTML content, create a single slide
        if not updated_slides_data and slides_html.strip():
            slide_data = {
                "page_number": 1,
                "title": "编辑后的Slide",
                "html_content": slides_html,
                "is_user_edited": True,
            }
            updated_slides_data.append(slide_data)

        logger.info(f"Extracted {len(updated_slides_data)} slides from HTML content")
        return updated_slides_data

    except Exception as e:
        logger.error(f"Error extracting slides from HTML: {e}")
        # Fall back to marking existing slides as edited
        if existing_slides_data:
            updated_slides_data = []
            for slide in existing_slides_data:
                slide_copy = slide.copy()
                slide_copy["is_user_edited"] = True
                updated_slides_data.append(slide_copy)
            return updated_slides_data
        else:
            return []


@router.get("/home", response_class=HTMLResponse)
async def web_home(request: Request, user: Optional[User] = Depends(get_current_user_optional)):
    """Public home page - always show landing, do not redirect for logged-in users.
    Also attach user to request.state for correct nav rendering."""
    if user and not getattr(request.state, "user", None):
        # Make sure templates can read the login state via request.state.user
        try:
            request.state.user = user
        except Exception:
            pass

    return templates.TemplateResponse(
        "index.html",
        {
            "request": request,
            "ai_provider": ai_config.default_ai_provider,
            "available_providers": ai_config.get_available_providers(),
            # Pass user explicitly as well for robustness
            "user": user.to_dict() if user else None,
        },
    )


@router.get("/ai-config", response_class=HTMLResponse)
async def web_ai_config(request: Request, user: User = Depends(get_current_user_required)):
    """AI configuration page"""
    from ..services.config_service import get_config_service

    config_service = get_config_service()
    current_config = config_service.get_all_config()

    # Temporarily suppress info/debug logs for this module while rendering this page
    level_filter = _LevelFilter(level=logging.WARNING)
    module_logger = logging.getLogger(__name__)
    try:
        module_logger.addFilter(level_filter)
        return templates.TemplateResponse(
            "ai_config.html",
            {
                "request": request,
                "current_provider": ai_config.default_ai_provider,
                "available_providers": ai_config.get_available_providers(),
                # build provider status separately to avoid long inline comprehension
                "provider_status": _build_provider_status(ai_config),
                "current_config": current_config,
                        "user": user.to_dict(),
                        "is_admin": getattr(user, "is_admin", False),
            },
        )
    finally:
        try:
            module_logger.removeFilter(level_filter)
        except Exception:
            pass


@router.post("/api/ai/providers/openai/models")
async def get_openai_models(request: Request, user: User = Depends(get_current_user_required)):
    """Proxy endpoint to get OpenAI models list, avoiding CORS issues - uses frontend provided config"""
    try:

        import aiohttp

        # Get configuration from frontend request
        data = await request.json()
        base_url = data.get("base_url", "https://api.openai.com/v1")
        api_key = data.get("api_key", "")

        base_info = f"Frontend requested models from: {base_url}"
        logger.info(base_info)

        if not api_key:
            return {"success": False, "error": "API Key is required"}

        # Build models URL safely (ensure /v1)
        models_url = build_api_url(base_url, "models", ensure_v1=True)
        logger.info(f"Fetching models from: {models_url}")

        # Make request to OpenAI API using frontend provided credentials
        async with aiohttp.ClientSession() as session:
            headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            }
            # Some upstream proxies expect an X-Api-Key header instead of or in addition to Bearer
            headers["X-Api-Key"] = api_key

            async with session.get(models_url, headers=headers, timeout=30) as response:
                if response.status == 200:
                    data = await response.json()

                    # Filter and sort models
                    models = []
                    if "data" in data and isinstance(data["data"], list):
                        for model in data["data"]:
                            if model.get("id"):
                                models.append(
                                    {
                                        "id": model["id"],
                                        "created": model.get("created", 0),
                                        "owned_by": model.get("owned_by", "unknown"),
                                    }
                                )

                        # Sort models with GPT-4 first, then GPT-3.5, then others
                        def get_priority(model_id):
                            if "gpt-4" in model_id:
                                return 0
                            elif "gpt-3.5" in model_id:
                                return 1
                            else:
                                return 2

                        models.sort(key=lambda x: (get_priority(x["id"]), x["id"]))
                    logger.info(f"Successfully fetched {len(models)} models from {base_url}")
                    return {"success": True, "models": models}
                else:
                    error_text = await response.text()
                    logger.error(
                        f"Failed to fetch models from {base_url}: {response.status} - {error_text}"
                    )
                    return {
                        "success": False,
                        "error": f"API returned status {response.status}: {error_text}",
                    }

    except Exception as e:
        logger.error(f"Error fetching OpenAI models from frontend config: {e}")
        return {"success": False, "error": str(e)}


@router.post("/api/ai/providers/anthropic/models")
async def get_anthropic_models(request: Request, user: User = Depends(get_current_user_required)):
    """Proxy endpoint to get Anthropic models list using frontend-provided config."""
    try:
        import aiohttp

        # Parse body defensively to avoid FastAPI returning 400 on malformed JSON.
        try:
            data = await request.json()
            if not isinstance(data, dict):
                data = {}
        except Exception as parse_err:
            # Attempt to read raw body for logging, but continue with empty payload
            try:
                raw = await request.body()
                raw_text = raw.decode("utf-8", errors="ignore")
            except Exception:
                raw_text = None
            logger.debug(
                "Google provider test: failed to parse JSON body: %s; raw=%s",
                str(parse_err),
                _sanitize_text(raw_text) if raw_text else "(no body)",
            )
            data = {}
        base_url = data.get("base_url", "https://api.anthropic.com").rstrip("/")
        api_key = data.get("api_key", "")
        version = data.get("api_version", "2023-06-01")

        if not api_key:
            return {"success": False, "error": "API Key is required"}
        # Anthropic models API endpoint (append v1/models)
        url = build_api_url(base_url, "v1/models")

        async with aiohttp.ClientSession() as session:
            headers = {
                "x-api-key": api_key,
                "anthropic-version": version,
                "content-type": "application/json",
            }
            async with session.get(url, headers=headers, timeout=30) as resp:
                text = await resp.text()
                if resp.status != 200:
                    logger.error("Anthropic models fetch failed %s: %s", resp.status, text)
                    err = f"HTTP {resp.status}: {text}"
                    return {"success": False, "error": err}
                try:
                    data = await resp.json()
                except Exception:
                    return {"success": False, "error": text}

                models = []
                # Anthropic API returns { "data": [ {"id": "claude-3-5-sonnet-20241022", ...}, ... ] }
                if isinstance(data, dict):
                    items = data.get("data", [])
                    for item in items:
                        model_id = item.get("id")
                        if model_id:
                            models.append({"id": model_id})

                return {"success": True, "models": models}
    except Exception as e:
        logger.error(f"Error fetching Anthropic models: {e}")
        return {"success": False, "error": str(e)}


@router.post("/api/ai/providers/google/models")
async def get_google_models(request: Request, user: User = Depends(get_current_user_required)):
    """Proxy endpoint to get Google Gemini models list using frontend-provided config."""
    try:
        import aiohttp

        body = await request.json()
        base_url = body.get("base_url", "https://generativelanguage.googleapis.com").rstrip("/")
        api_key = body.get("api_key", "")
        if not api_key:
            return {"success": False, "error": "API Key is required"}

        # Use v1beta endpoint which returns more models
        models_endpoint = build_api_url(base_url, "v1beta/models")
        url = models_endpoint + "?key=" + api_key

        logger.info("Calling Google v1beta API: %s", models_endpoint + "?key=" + _redact_key(api_key))

        async with aiohttp.ClientSession() as session:
            async with session.get(url, timeout=30) as resp:
                text = await resp.text()
                logger.info(f"Google v1beta API response status: {resp.status}")
                logger.info(f"Google v1beta API response length: {len(text)} characters")

                if resp.status != 200:
                    safe_text = _sanitize_text(text)
                    logger.error(f"Google v1beta models fetch failed {resp.status}: {safe_text}")
                    return {"success": False, "error": f"HTTP {resp.status}: {safe_text}"}

                try:
                    data = await resp.json()
                    structure = list(data.keys()) if isinstance(data, dict) else type(data)
                    logger.info("Google v1beta API response structure: %s", structure)
                except Exception:
                    logger.error(
                        f"Failed to parse Google v1beta API response as JSON: {_sanitize_text(text)}"
                    )
                    return {"success": False, "error": _sanitize_text(text)}

                models = []
                # Google API returns {"models": [{"name": "models/gemini-pro", ...}, ...]}
                items = data.get("models", [])
                logger.info(f"Google v1beta API returned {len(items)} models")

                for item in items:
                    name = item.get("name", "")
                    if name and name.startswith("models/"):
                        # Extract model name from "models/gemini-pro" format
                        model_id = name.replace("models/", "")
                        # Include all available models
                        models.append({"id": model_id})
                        logger.debug(f"Added model: {model_id}")

                logger.info(f"Returning {len(models)} models to frontend")
                return {"success": True, "models": models}
    except Exception as e:
        logger.error(f"Error fetching Google models: {e}")
        return {"success": False, "error": "Failed to fetch Google models"}


@router.post("/api/ai/providers/azure_openai/models")
async def get_azure_openai_deployments(
    request: Request, user: User = Depends(get_current_user_required)
):
    """Proxy endpoint to list Azure OpenAI deployments (used as model names)."""
    try:
        import aiohttp

        body = await request.json()
        endpoint = (body.get("endpoint") or body.get("base_url") or "").rstrip("/")
        api_key = body.get("api_key", "")
        api_version = body.get("api_version", "2024-02-15-preview")
        if not endpoint:
            return {"success": False, "error": "Endpoint is required"}
        if not api_key:
            return {"success": False, "error": "API Key is required"}
        deployments_endpoint = build_api_url(endpoint, "openai/deployments")
        url = deployments_endpoint + "?api-version=" + api_version

        async with aiohttp.ClientSession() as session:
            headers = {"api-key": api_key}
            async with session.get(url, headers=headers, timeout=30) as resp:
                text = await resp.text()
                if resp.status != 200:
                    logger.error(f"Azure deployments fetch failed {resp.status}: {text}")
                    return {"success": False, "error": f"HTTP {resp.status}: {text}"}
                try:
                    data = await resp.json()
                except Exception:
                    return {"success": False, "error": text}

                models = []
                # Could be {"data": [...]} or {"value": [...]} depending on API
                items = data.get("data") or data.get("value") or []
                for it in items:
                    name = it.get("name") or it.get("id")
                    if name:
                        models.append({"id": name})
                return {"success": True, "models": models}
    except Exception as e:
        logger.error(f"Error fetching Azure OpenAI deployments: {e}")
        return {"success": False, "error": str(e)}


@router.post("/api/ai/providers/ollama/models")
async def get_ollama_models(request: Request, user: User = Depends(get_current_user_required)):
    """Proxy endpoint to list Ollama local models (tags)."""
    try:
        import aiohttp
        from ..services.config_service import get_config_service

        body = await request.json()
        base_url = body.get("base_url", "http://localhost:11434").rstrip("/")
        url = build_api_url(base_url, "api/tags")

        # Read optional upstream API key for Ollama (kept in server-side config)
        cfg = get_config_service().get_config_by_category("ai_providers")
        ollama_key = cfg.get("ollama_api_key") if cfg else None

        headers = {}
        if ollama_key:
            headers["Authorization"] = f"Bearer {ollama_key}"

        async with aiohttp.ClientSession() as session:
            async with session.get(url, headers=headers or None, timeout=15) as resp:
                text = await resp.text()
                if resp.status != 200:
                    return {"success": False, "error": f"HTTP {resp.status}: {_sanitize_text(text)}"}
                try:
                    data = await resp.json()
                except Exception:
                    return {"success": False, "error": _sanitize_text(text)}

                models = []
                for m in data.get("models", []):
                    name = m.get("name") or m.get("model")
                    if name:
                        models.append({"id": name})

                return {"success": True, "models": models}
    except Exception as e:
        logger.error(f"Error fetching Ollama models: {e}")
    return {"success": False, "error": "Failed to fetch Ollama models"}


@router.post("/api/ai/providers/openai/validate")
async def validate_openai_api_key(
    request: Request, user: User = Depends(get_current_user_required)
):
    """Validate OpenAI API Key"""
    try:
        import aiohttp

        # Get configuration from frontend request
        data = await request.json()
        base_url = data.get("base_url", "https://api.openai.com/v1")
        api_key = data.get("api_key", "")

        logger.info("Validating API Key for: %s", base_url)

        if not api_key:
            return {"success": False, "error": "API Key is required"}
        # Build models URL safely (ensure /v1)
        models_url = build_api_url(base_url, "models", ensure_v1=True)

        async with aiohttp.ClientSession() as session:
            headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            }
            # Some proxies require X-Api-Key header as well
            headers["X-Api-Key"] = api_key

            async with session.get(models_url, headers=headers, timeout=10) as response:
                if response.status == 200:
                    data = await response.json()
                    model_count = len(data.get("data", []))
                    return {
                        "success": True,
                        "message": "API Key 有效，可访问 %d 个模型" % model_count,
                        "model_count": model_count,
                    }
                else:
                    error_text = await response.text()
                    logger.error(f"API Key validation failed: {response.status} - {error_text}")

                    # 解析错误信息
                    try:
                        error_data = (
                            await response.json()
                            if response.content_type == "application/json"
                            else {}
                        )
                        error_msg = error_data.get("error", {}).get("message", error_text)
                    except Exception:
                        error_msg = error_text

                    return {
                        "success": False,
                        "error": f"API Key 验证失败 (状态码: {response.status}): {error_msg}",
                    }

    except Exception as e:
        logger.error(f"Error validating OpenAI API Key: {e}")
        return {"success": False, "error": f"验证过程出错: {str(e)}"}


@router.post("/api/ai/providers/openai/test")
async def test_openai_provider_proxy(
    request: Request, user: User = Depends(get_current_user_required)
):
    """Proxy endpoint to test OpenAI provider, avoiding CORS issues - uses frontend provided config"""
    try:
        import aiohttp

        # Get configuration from frontend request
        data = await request.json()
        base_url = data.get("base_url", "https://api.openai.com/v1")
        api_key = data.get("api_key", "")
        model = data.get("model", "gpt-4o")

        logger.info("Frontend requested test with base_url=%s model=%s", base_url, model)

        if not api_key:
            return {"success": False, "error": "API Key is required"}
        # Build chat URL safely (ensure /v1)
        chat_url = build_api_url(base_url, "chat/completions", ensure_v1=True)
        logger.info(f"Testing OpenAI provider at: {chat_url}")

        # Make test request to OpenAI API using frontend provided credentials
        async with aiohttp.ClientSession() as session:
            headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            }
            # Some proxies require X-Api-Key header as well
            headers["X-Api-Key"] = api_key

            payload = {
                "model": model,
                "messages": [
                    {
                        "role": "user",
                        "content": "Say 'Hello, I am working!' in exactly 5 words.",
                    }
                ],
                "max_tokens": 20,
                "temperature": 0,
            }

            try:
                async with session.post(chat_url, headers=headers, json=payload, timeout=30) as response:
                    resp_text = await response.text()
                    if response.status == 200:
                        try:
                            data = json.loads(resp_text)
                        except Exception:
                            data = None

                        logger.info(f"Test successful for {base_url} with model {model}")

                        response_preview = None
                        try:
                            if isinstance(data, dict):
                                choices = data.get("choices")
                                if isinstance(choices, list) and len(choices) > 0:
                                    first = choices[0]
                                    if isinstance(first, dict):
                                        message = first.get("message")
                                        if isinstance(message, dict):
                                            response_preview = message.get("content")
                        except Exception:
                            response_preview = None

                        if not response_preview:
                            response_preview = str(data)[:500] if data is not None else resp_text[:500]

                        usage = data.get("usage") if isinstance(data, dict) else None
                        if not usage:
                            usage = {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}

                        return {
                            "success": True,
                            "status": "success",
                            "provider": "openai",
                            "model": model,
                            "response_preview": response_preview,
                            "usage": usage,
                        }
                    else:
                        try:
                            error_data = json.loads(resp_text)
                            error_message = error_data.get("error", {}).get("message") or str(error_data)
                        except Exception:
                            error_message = f"API returned status {response.status}: {resp_text}"

                        logger.error(f"Test failed for {base_url}: {error_message}")
                        return {"success": False, "status": "error", "error": error_message}
            except Exception as aio_err:
                logger.error(f"Exception during OpenAI test request: {aio_err}")
                return {"success": False, "status": "error", "error": str(aio_err)}

    except Exception as e:
        logger.error(f"Error testing OpenAI provider with frontend config: {e}")
        return {"success": False, "status": "error", "error": str(e)}


@router.post("/api/ai/providers/google/test")
async def test_google_provider_proxy(request: Request):
    """Backend proxy to test Google Gemini provider using frontend-provided config.

    This avoids CORS and supports both official Google endpoint (query param ?key=)
    and custom gateways that require Authorization/X-Api-Key headers.
    """
    try:
        import aiohttp

        # Parse body defensively to avoid FastAPI returning 400 on malformed JSON.
        try:
            data = await request.json()
            if not isinstance(data, dict):
                data = {}
        except Exception as parse_err:
            # Attempt to read raw body for logging, but continue with empty payload
            try:
                raw = await request.body()
                raw_text = raw.decode("utf-8", errors="ignore")
            except Exception:
                raw_text = None
            logger.debug(
                "Google provider test: failed to parse JSON body: %s; raw=%s",
                str(parse_err),
                _sanitize_text(raw_text) if raw_text else "(no body)",
            )
            data = {}

        base_url = (data.get("base_url") or "https://generativelanguage.googleapis.com").rstrip("/")
        api_key = (data.get("api_key") or "").strip()
        model = (data.get("model") or "gemini-1.5-flash").strip()

        if not api_key:
            return {"success": False, "status": "error", "error": "API Key is required"}

        # Build REST URL: /v1beta/models/{model}:generateContent
        # Use helper to safely join paths
        gen_url = build_api_url(base_url, "v1beta/models", f"{model}:generateContent")

        # Some gateways only accept headers; Google's official accepts query param ?key=
        use_query_key = base_url.endswith("generativelanguage.googleapis.com")
        if use_query_key:
            gen_url = f"{gen_url}?key={api_key}"

        payload = {
            "contents": [
                {
                    "parts": [
                        {"text": 'Say "Hello, I am working!" in exactly 5 words.'}
                    ]
                }
            ],
            "generationConfig": {"maxOutputTokens": 20, "temperature": 0},
        }

        headers = {"Content-Type": "application/json"}
        # Provide multiple header styles for better compatibility with proxies
        headers["Authorization"] = f"Bearer {api_key}"
        headers["X-Api-Key"] = api_key
        headers["x-goog-api-key"] = api_key

        async with aiohttp.ClientSession() as session:
            async with session.post(gen_url, json=payload, headers=headers, timeout=30) as resp:
                text = await resp.text()
                if resp.status == 200:
                    try:
                        data = await resp.json()
                    except Exception:
                        # Return raw text if JSON parse fails
                        data = None

                    # Extract a short preview
                    preview = None
                    try:
                        if isinstance(data, dict):
                            cand = data.get("candidates")
                            if isinstance(cand, list) and cand:
                                c0 = cand[0]
                                content = (
                                    (c0 or {}).get("content") if isinstance(c0, dict) else None
                                )
                                parts = (content or {}).get("parts") if isinstance(content, dict) else None
                                if isinstance(parts, list) and parts:
                                    preview = (parts[0] or {}).get("text")
                    except Exception:
                        preview = None

                    if not preview:
                        preview = (text or "")[:500]

                    usage = None
                    if isinstance(data, dict):
                        um = data.get("usageMetadata")
                        if isinstance(um, dict):
                            usage = {
                                "prompt_tokens": um.get("promptTokenCount", 0),
                                "completion_tokens": um.get("candidatesTokenCount", 0),
                                "total_tokens": um.get("totalTokenCount", 0),
                            }
                    if not usage:
                        usage = {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}

                    return {
                        "success": True,
                        "status": "success",
                        "provider": "google",
                        "model": model,
                        "response_preview": preview,
                        "usage": usage,
                    }
                else:
                    # Try parse JSON error
                    err_msg = text
                    try:
                        err_json = json.loads(text)
                        err_msg = (
                            err_json.get("error", {}).get("message")
                            or err_json.get("message")
                            or str(err_json)
                            or err_msg
                        )
                    except Exception:
                        pass
                    return {"success": False, "status": "error", "error": f"HTTP {resp.status}: {err_msg}"}

    except Exception as e:
        logger.error(f"Error testing Google provider with frontend config: {e}")
        return {"success": False, "status": "error", "error": str(e)}

@router.post("/api/ai/providers/ollama/test")
async def test_ollama_provider_proxy(request: Request):
    """Backend proxy to test Ollama provider (validates base_url and shields credentials/CORS)."""
    try:
        import aiohttp
        # Parse body defensively
        try:
            data = await request.json()
            if not isinstance(data, dict):
                data = {}
        except Exception:
            data = {}
        logger.info(f"Ollama test payload: {_sanitize_dict(data) if isinstance(data, dict) else '(invalid payload)'}")
        base_url = (data.get("base_url") or "http://localhost:11434").rstrip("/")
        model = (data.get("model") or "llama2").strip()

        # Basic validation of URL scheme
        parsed = urllib.parse.urlparse(base_url)
        if parsed.scheme not in ("http", "https"):
            return JSONResponse({"success": False, "status": "error", "error": "Invalid base_url scheme"}, status_code=200)
        gen_url = build_api_url(base_url, "api/generate")

        payload = {
            "model": model,
            "prompt": 'Say "Hello, I am working!" in exactly 5 words.',
            "stream": False,
            "options": {"num_predict": 20, "temperature": 0},
        }

        # Only use API key passed from body; avoid touching server-side config here
        use_key = (data.get("api_key") or "").strip() or None

        headers = {}
        if use_key:
            headers["Authorization"] = f"Bearer {use_key}"
            headers["X-Api-Key"] = use_key

        async with aiohttp.ClientSession() as session:
            # 1) Quick ping Ollama (/api/tags) to avoid 500s when service is down
            try:
                tags_url = build_api_url(base_url, "api/tags")
                async with session.get(tags_url, headers=headers or None, timeout=10) as ping:
                    ping_text = await ping.text()
                    if ping.status != 200:
                        return JSONResponse({
                            "success": False,
                            "status": "error",
                            "provider": "ollama",
                            "error": f"无法连接到 Ollama ({ping.status})",
                            "detail": ping_text[:500]
                        }, status_code=200)

                    # Validate model presence if parsable
                    try:
                        tags_json = json.loads(ping_text)
                        if isinstance(tags_json, dict) and model:
                            names = [m.get("name") or m.get("model") for m in tags_json.get("models", [])]
                            if names and model not in names:
                                return JSONResponse({
                                    "success": False,
                                    "status": "error",
                                    "provider": "ollama",
                                    "error": f"模型未找到: {model}",
                                    "detail": f"已安装模型: {', '.join([n for n in names if n])}"
                                }, status_code=200)
                    except Exception:
                        pass
            except Exception as ping_err:
                logger.info(f"Ollama ping failed: {ping_err}")
                return JSONResponse({
                    "success": False,
                    "status": "error",
                    "provider": "ollama",
                    "error": "Ollama 服务未运行或无法连接",
                    "detail": f"请确保服务可通过 {base_url} 访问，并已拉取模型 {model}"
                }, status_code=200)

            # 2) Generate test
            try:
                safe_url = gen_url  # URL contains only local host/port and path
                logger.info(f"Calling Ollama generate URL: {safe_url} with model={model}")
                async with session.post(gen_url, json=payload, headers=headers or None, timeout=30) as resp:
                    text = await resp.text()
                    if resp.status == 200:
                        try:
                            resp_json = json.loads(text)
                        except Exception:
                            resp_json = None

                        response_preview = None
                        if isinstance(resp_json, dict):
                            response_preview = resp_json.get("response") or resp_json.get("output")

                        if not response_preview:
                            response_preview = text[:500]

                        return JSONResponse({
                            "success": True,
                            "status": "success",
                            "provider": "ollama",
                            "model": model,
                            "response_preview": response_preview,
                        }, status_code=200)
                    else:
                        try:
                            err = json.loads(text)
                            err_msg = err.get("error") or str(err)
                        except Exception:
                            err_msg = f"HTTP {resp.status}: {_sanitize_text(text)}"
                        logger.error(f"Ollama test failed for {safe_url}: {_sanitize_text(err_msg)}")
                        return JSONResponse({"success": False, "status": "error", "error": _sanitize_text(err_msg)}, status_code=200)
            except Exception as aio_err:
                logger.exception(f"Exception during Ollama test request: {aio_err}")
                return JSONResponse({"success": False, "status": "error", "error": _sanitize_text(str(aio_err))}, status_code=200)

    except Exception as e:
        logger.exception(f"Error testing Ollama provider: {e}")
        return JSONResponse({"success": False, "status": "error", "error": _sanitize_text(str(e))}, status_code=200)


@router.get("/scenarios", response_class=HTMLResponse)
async def web_scenarios(
    request: Request, user: Optional[User] = Depends(get_current_user_optional)
):
    """Scenarios selection page"""
    scenarios = [
        {
            "id": "general",
            "name": "通用",
            "description": "适用于各种通用场景的Slide模板",
            "icon": "📋",
        },
        {
            "id": "tourism",
            "name": "旅游观光",
            "description": "旅游线路、景点介绍等旅游相关Slide",
            "icon": "🌍",
        },
        {
            "id": "education",
            "name": "儿童科普",
            "description": "适合儿童的科普教育Slide",
            "icon": "🎓",
        },
        {
            "id": "analysis",
            "name": "深入分析",
            "description": "数据分析、研究报告等深度分析Slide",
            "icon": "📊",
        },
        {
            "id": "history",
            "name": "历史文化",
            "description": "历史事件、文化介绍等人文类Slide",
            "icon": "🏛️",
        },
        {
            "id": "technology",
            "name": "科技技术",
            "description": "技术介绍、产品发布等科技类Slide",
            "icon": "💻",
        },
        {
            "id": "business",
            "name": "方案汇报",
            "description": "商业计划、项目汇报等商务Slide",
            "icon": "💼",
        },
    ]
    return templates.TemplateResponse(
        "scenarios.html", {"request": request, "scenarios": scenarios}
    )


# Legacy route removed - now using /projects/create for new project workflow

# Legacy task status route removed - now using project detail pages

# Legacy preview route removed - now using project-based preview at /projects/{project_id}/fullscreen

# Legacy tasks list route removed - now using /projects for project management


@router.post("/upload", response_class=HTMLResponse)
async def web_upload_file(
    request: Request,
    file: UploadFile = File(...),
    user: User = Depends(get_current_user_required),
):
    """Upload file via web interface"""
    try:
        # Validate file type
        allowed_types = [".docx", ".pdf", ".txt", ".md"]
        file_extension = "." + file.filename.split(".")[-1].lower()

        if file_extension not in allowed_types:
            return templates.TemplateResponse(
                "upload_result.html",
                {
                    "request": request,
                    "success": False,
                    "error": f"Unsupported file type. Allowed types: {', '.join(allowed_types)}",
                },
            )

        # Read file content in thread pool to avoid blocking
        content = await file.read()

        # Process file in thread pool
        processed_content = await ppt_service.process_uploaded_file(
            filename=file.filename, content=content, file_type=file_extension
        )

        return templates.TemplateResponse(
            "upload_result.html",
            {
                "request": request,
                "success": True,
                "filename": file.filename,
                "size": len(content),
                "type": file_extension,
                "processed_content": (
                    processed_content[:500] + "..."
                    if len(processed_content) > 500
                    else processed_content
                ),
            },
        )

    except Exception as e:
        return templates.TemplateResponse(
            "upload_result.html",
            {"request": request, "success": False, "error": str(e)},
        )


@router.get("/demo", response_class=HTMLResponse)
async def web_demo(request: Request, user: User = Depends(get_current_user_required)):
    """Demo page with sample Slide"""
    # Create a demo Slide
    demo_request = PPTGenerationRequest(
        scenario="technology",
        topic="人工智能技术发展趋势",
        requirements="面向技术人员的深度分析",
        network_mode=False,
        language="zh",
    )

    task_id = "demo-" + str(uuid.uuid4())[:8]
    result = await ppt_service.generate_ppt(task_id, demo_request)

    return templates.TemplateResponse(
        "demo.html",
        {
            "request": request,
            "task_id": task_id,
            "outline": result.get("outline"),
            "slides_html": result.get("slides_html"),
            "demo_topic": demo_request.topic,
        },
    )


@router.get("/research", response_class=HTMLResponse)
async def web_research_status(request: Request, user: User = Depends(get_current_user_required)):
    """DEEP Research status and management page"""
    return templates.TemplateResponse("research_status.html", {"request": request})


# New Project Management Routes


@router.get("/dashboard", response_class=HTMLResponse)
async def web_dashboard(request: Request, user: User = Depends(get_current_user_required)):
    """Project dashboard with overview"""
    try:
        # Get project statistics (RBAC: non-admin only sees own projects)
        owner_filter = None if user.is_admin else user.id
        # Load a reasonable page size for dashboard to avoid large payloads
        projects_response = await ppt_service.project_manager.list_projects(
            page=1, page_size=20, owner_id=owner_filter
        )
        projects = projects_response.projects

        total_projects = len(projects)
        completed_projects = len([p for p in projects if p.status == "completed"])
        in_progress_projects = len([p for p in projects if p.status == "in_progress"])
        draft_projects = len([p for p in projects if p.status == "draft"])

        # Get recent projects (last 5)
        recent_projects = sorted(projects, key=lambda x: x.updated_at, reverse=True)[:5]

        # Get active TODO boards (use already-loaded todo_board to avoid N+1 DB calls)
        active_todo_boards = []
        for project in projects:
            if project.status == "in_progress" and getattr(project, "todo_board", None):
                # project.todo_board is eager-loaded by the repository; reuse it directly
                active_todo_boards.append(project.todo_board)

        return templates.TemplateResponse(
            "project_dashboard.html",
            {
                "request": request,
                "total_projects": total_projects,
                "completed_projects": completed_projects,
                "in_progress_projects": in_progress_projects,
                "draft_projects": draft_projects,
                "recent_projects": recent_projects,
                "active_todo_boards": active_todo_boards[:3],  # Show max 3 boards
            },
        )

    except Exception as e:
        return templates.TemplateResponse("error.html", {"request": request, "error": str(e)})


@router.get("/projects", response_class=HTMLResponse)
async def web_projects_list(
    request: Request,
    page: int = 1,
    status: str = None,
    user: User = Depends(get_current_user_required),
    db: Session = Depends(get_db),
):
    """List all projects"""
    try:
        owner_filter = None if user.is_admin else user.id
        projects_response = await ppt_service.project_manager.list_projects(
            page=page, page_size=10, status=status, owner_id=owner_filter
        )

        # Build owner username map for admins
        owner_name_map = {}
        try:
            if getattr(user, "is_admin", False):
                owner_ids = list(
                    {
                        p.owner_id
                        for p in projects_response.projects
                        if getattr(p, "owner_id", None) is not None
                    }
                )
                if owner_ids:
                    rows = db.query(User).filter(User.id.in_(owner_ids)).all()
                    owner_name_map = {u.id: (u.username or f"user-{u.id}") for u in rows}
        except Exception:
            owner_name_map = {}

        return templates.TemplateResponse(
            "projects_list.html",
            {
                "request": request,
                "projects": projects_response.projects,
                "total": projects_response.total,
                "page": projects_response.page,
                "page_size": projects_response.page_size,
                "status_filter": status,
                "owner_name_map": owner_name_map,
            },
        )

    except Exception as e:
        return templates.TemplateResponse("error.html", {"request": request, "error": str(e)})


@router.get("/projects/{project_id}", response_class=HTMLResponse)
async def web_project_detail(
    request: Request, project_id: str, user: User = Depends(get_current_user_required)
):
    """Project detail page"""
    try:
        project = await ppt_service.project_manager.get_project(project_id)
        if not project:
            return templates.TemplateResponse(
                "error.html", {"request": request, "error": "Project not found"}
            )
        # RBAC: non-admin cannot access others' projects
        if not _user_can_access_project(user, project):
            return templates.TemplateResponse(
                "error.html", {"request": request, "error": "Access denied"}
            )

        todo_board = await ppt_service.get_project_todo_board(project_id)
        versions = await ppt_service.project_manager.get_project_versions(project_id)

        return templates.TemplateResponse(
            "project_detail.html",
            {
                "request": request,
                "project": project,
                "todo_board": todo_board,
                "versions": versions,
            },
        )

    except Exception as e:
        return templates.TemplateResponse("error.html", {"request": request, "error": str(e)})


@router.get("/projects/{project_id}/todo", response_class=HTMLResponse)
async def web_project_todo_board(
    request: Request, project_id: str, user: User = Depends(get_current_user_required)
):
    """TODO board page for a project with integrated editor"""
    try:
        # Validate project_id format (should be UUID-like)
        if project_id in [
            "template-selection",
            "todo",
            "edit",
            "preview",
            "fullscreen",
        ]:
            error_msg = f"无效的项目ID: {project_id}。\n\n"
            error_msg += "可能的原因：\n"
            error_msg += "1. URL格式错误，正确格式应为: /projects/[项目ID]/todo\n"
            error_msg += "2. 您可能访问了错误的链接\n\n"
            error_msg += "建议解决方案：\n"
            error_msg += "• 返回项目列表页面选择正确的项目\n"
            error_msg += "• 检查浏览器地址栏中的URL是否完整"

            return templates.TemplateResponse(
                "error.html", {"request": request, "error": error_msg}
            )

        # Check if project exists first
        project = await ppt_service.project_manager.get_project(project_id)
        if not project:
            return templates.TemplateResponse(
                "error.html",
                {
                    "request": request,
                    "error": f"项目不存在 (ID: {project_id})。请检查项目ID是否正确。",
                },
            )
        if not _user_can_access_project(user, project):
            return templates.TemplateResponse(
                "error.html",
                {"request": request, "error": "无权限访问该项目"},
            )

        todo_board = await ppt_service.get_project_todo_board(project_id)
        if not todo_board:
            return templates.TemplateResponse(
                "error.html",
                {
                    "request": request,
                    "error": f"项目 '{project.topic}' 的TODO看板不存在。请联系技术支持。",
                },
            )

        # Check if we should use the integrated editor version
        project = await ppt_service.project_manager.get_project(project_id)
        use_integrated_editor = (
            project
            and project.confirmed_requirements
            and len(todo_board.stages) > 2
            and (
                todo_board.stages[1].status in ["running", "completed"]
                or todo_board.stages[2].status in ["running", "completed"]
            )
        )

        # Also use integrated editor if Slide creation stage is about to start or running
        if (
            project
            and project.confirmed_requirements
            and len(todo_board.stages) > 2
            and todo_board.stages[1].status == "completed"
        ):
            use_integrated_editor = True

        template_name = (
            "todo_board_with_editor.html" if use_integrated_editor else "todo_board.html"
        )

        # Ensure project is not None for template
        template_context = {
            "request": request,
            "todo_board": todo_board,
            "aspect_ratio_settings": get_aspect_ratio_settings(),
        }

        # Only add project if it exists
        if project:
            template_context["project"] = project

        return templates.TemplateResponse(template_name, template_context)

    except Exception as e:
        return templates.TemplateResponse("error.html", {"request": request, "error": str(e)})


@router.get("/projects/{project_id}/fullscreen", response_class=HTMLResponse)
async def web_project_fullscreen(
    request: Request, project_id: str, user: User = Depends(get_current_user_required)
):
    """Fullscreen preview of project Slide with modern presentation interface"""
    try:
        # 直接从数据库获取最新的项目数据，确保数据实时性
        from ..services.db_project_manager import DatabaseProjectManager

        db_manager = DatabaseProjectManager()
        project = await db_manager.get_project(project_id)

        if not project:
            return templates.TemplateResponse(
                "error.html", {"request": request, "error": "项目未找到"}
            )
        if not _user_can_access_project(user, project):
            return templates.TemplateResponse(
                "error.html", {"request": request, "error": "无权限访问该项目"}
            )

        # 检查是否有幻灯片数据
        if not project.slides_data or len(project.slides_data) == 0:
            return templates.TemplateResponse(
                "error.html", {"request": request, "error": "Slide尚未生成或无幻灯片内容"}
            )

        # 使用新的分享演示模板
        return templates.TemplateResponse(
            "project_fullscreen_presentation.html",
            {
                "request": request,
                "project": project,
                "slides_count": len(project.slides_data),
            },
        )

    except Exception as e:
        logger.error(f"Error in fullscreen presentation: {e}")
        return templates.TemplateResponse(
            "error.html", {"request": request, "error": f"加载演示时出错: {str(e)}"}
        )


@router.get("/api/projects/{project_id}/slides-data")
async def get_project_slides_data(project_id: str, user: User = Depends(get_current_user_required)):
    """获取项目最新的幻灯片数据 - 用于分享演示实时更新"""
    try:
        # 直接从数据库获取最新数据
        from ..services.db_project_manager import DatabaseProjectManager

        db_manager = DatabaseProjectManager()
        project = await db_manager.get_project(project_id)

        if not project:
            raise HTTPException(status_code=404, detail="项目未找到")
        if not _user_can_access_project(user, project):
            raise HTTPException(status_code=403, detail="无权限访问该项目")

        if not project.slides_data or len(project.slides_data) == 0:
            return {
                "status": "no_slides",
                "message": "Slide尚未生成",
                "slides_data": [],
                "total_slides": 0,
            }

        # Compute per-slide content hashes to allow clients to diff changes
        try:
            slides_meta = []
            for i, s in enumerate(project.slides_data):
                content = (s.get('html_content') or '')
                h = hashlib.sha256(content.encode('utf-8')).hexdigest()
                slides_meta.append({
                    'index': i,
                    'hash': h,
                    'page_number': s.get('page_number', i + 1),
                })
        except Exception:
            slides_meta = []

        return {
            "status": "success",
            "slides_data": project.slides_data,
            "slides_meta": slides_meta,
            "total_slides": len(project.slides_data),
            "project_title": project.title,
            "updated_at": project.updated_at,
        }

    except Exception as e:
        logger.error(f"Error getting slides data: {e}")
        raise HTTPException(status_code=500, detail=f"获取幻灯片数据失败: {str(e)}")


@router.get("/test/slides-navigation", response_class=HTMLResponse)
async def test_slides_navigation(request: Request, user: User = Depends(get_current_user_required)):
    """测试幻灯片导航功能"""
    with open("test_slides_navigation.html", "r", encoding="utf-8") as f:
        content = f.read()
    return HTMLResponse(content=content)


@router.get("/temp/{file_path:path}")
async def serve_temp_file(file_path: str, user: User = Depends(get_current_user_required)):
    """Serve temporary slide files"""
    try:
        # Construct the full path to the temp file using system temp directory
        import tempfile

        temp_dir = Path(tempfile.gettempdir()) / "flowslide"
        full_path = temp_dir / file_path

        # Security check: ensure the file is within the temp directory
        if not str(full_path.resolve()).startswith(str(temp_dir.resolve())):
            raise HTTPException(status_code=403, detail="Access denied")

        # Check if file exists
        if not full_path.exists():
            raise HTTPException(status_code=404, detail="File not found")

        # Return the file
        return FileResponse(
            path=str(full_path),
            media_type="text/html; charset=utf-8",
            headers={"Cache-Control": "no-cache"},
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/projects/create", response_class=HTMLResponse)
async def web_create_project(
    request: Request,
    scenario: str = Form(...),
    topic: str = Form(...),
    requirements: str = Form(None),
    language: str = Form("zh"),
    network_mode: bool = Form(False),
    existing_report: str = Form(None),
    report_file: UploadFile | None = File(None),
    user: User = Depends(get_current_user_required),
):
    """Create new project via web interface"""
    try:
        # Create project request
        uploaded_content = None
        use_file_content = False
        # 1. 处理用户上传的研究报告文件（优先级高于选择现有）
        if report_file and report_file.filename:
            try:
                fname = report_file.filename
                if not any(fname.lower().endswith(ext) for ext in [".md", ".txt"]):
                    raise ValueError("仅支持 .md 或 .txt 研究报告文件")
                raw = await report_file.read()
                text = raw.decode("utf-8", errors="ignore")
                if len(text) > 120_000:
                    text = text[:120_000]
                uploaded_content = text
                use_file_content = True
                network_mode = False  # 使用本地内容不再联网
                tag = f"(使用上传研究报告: {fname})"
                requirements = (requirements + "\n" + tag) if requirements else tag
                # 将上传内容持久化写入 research_reports 目录（不覆盖已有，同 upload API 规则简化版）
                try:
                    from pathlib import Path
                    reports_dir = Path("research_reports"); reports_dir.mkdir(exist_ok=True)
                    import re, os, time as _t
                    safe = re.sub(r"[^A-Za-z0-9_.\-\u4e00-\u9fa5]", "_", fname)[:100]
                    if not safe.lower().endswith(('.md','.txt')): safe += '.md'
                    candidate = safe
                    if (reports_dir / candidate).exists():
                        stem, ext = os.path.splitext(candidate)
                        candidate = f"{stem}_{int(_t.time())}{ext}"
                    (reports_dir / candidate).write_text(text, encoding='utf-8')
                except Exception as _se:
                    logger.warning(f"保存上传研究报告副本失败: {_se}")
            except Exception as e:
                logger.warning(f"读取上传研究报告失败: {e}")
        # 2. 否则如果选择已有本地报告
        elif existing_report:
            from pathlib import Path
            reports_dir = Path("research_reports")
            safe_name = existing_report.replace("..", "").strip("/\\")
            target = reports_dir / safe_name
            if target.is_file() and target.suffix.lower() in {".md", ".txt"}:
                try:
                    uploaded_content = target.read_text(encoding="utf-8", errors="ignore")
                    # 限制大小防止过长
                    if len(uploaded_content) > 120_000:
                        uploaded_content = uploaded_content[:120_000]
                    use_file_content = True
                    # 如果选了已有报告，就不再联网
                    network_mode = False
                    if requirements:
                        requirements = requirements + f"\n(使用已有研究报告: {safe_name})"
                    else:
                        requirements = f"使用已有研究报告: {safe_name}"
                except Exception as e:
                    logger.warning(f"读取已有研究报告失败: {e}")

        project_request = PPTGenerationRequest(
            scenario=scenario,
            topic=topic,
            requirements=requirements,
            network_mode=network_mode,
            language=language,
            uploaded_content=uploaded_content,
            use_file_content=use_file_content,
        )

        # Create project with TODO board (without starting workflow yet)
        # Attach current user as owner to enforce RBAC
        project = await ppt_service.project_manager.create_project(
            project_request, owner_id=getattr(user, "id", None)
        )

        # Update project status to in_progress
        await ppt_service.project_manager.update_project_status(project.project_id, "in_progress")

        # Redirect directly to TODO page without showing redirect page
        from fastapi.responses import RedirectResponse

        return RedirectResponse(url=f"/projects/{project.project_id}/todo", status_code=302)

    except Exception as e:
        return templates.TemplateResponse("error.html", {"request": request, "error": str(e)})


@router.post("/projects/{project_id}/start-workflow")
async def start_project_workflow(project_id: str, user: User = Depends(get_current_user_required)):
    """Start the AI workflow for a project (only if requirements are confirmed)"""
    try:
        # Get project
        project = await ppt_service.project_manager.get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        if not _user_can_access_project(user, project):
            raise HTTPException(status_code=403, detail="Access denied")

        # Check if requirements are confirmed
        if not project.confirmed_requirements:
            return {
                "status": "waiting",
                "message": "Waiting for requirements confirmation",
            }

        # Extract network_mode from project metadata
        network_mode = False
        if project.project_metadata and isinstance(project.project_metadata, dict):
            network_mode = project.project_metadata.get("network_mode", False)

        # Create project request from project data
        confirmed_requirements = project.confirmed_requirements or {}
        project_request = PPTGenerationRequest(
            scenario=project.scenario,
            topic=project.topic,
            requirements=project.requirements,
            language="zh",  # Default language
            network_mode=network_mode,
            target_audience=confirmed_requirements.get("target_audience", "普通大众"),
            ppt_style=confirmed_requirements.get("ppt_style", "general"),
            custom_style_prompt=confirmed_requirements.get("custom_style_prompt"),
            description=confirmed_requirements.get("description"),
        )

        # Start the workflow in background
        asyncio.create_task(ppt_service._execute_project_workflow(project_id, project_request))

        return {"status": "success", "message": "Workflow started"}

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/projects/{project_id}/requirements", response_class=HTMLResponse)
async def project_requirements_page(
    request: Request, project_id: str, user: User = Depends(get_current_user_required)
):
    """Show project requirements confirmation page"""
    try:
        # Get project
        project = await ppt_service.project_manager.get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        if not _user_can_access_project(user, project):
            raise HTTPException(status_code=403, detail="Access denied")

        # 提供默认的展示类型选项，不再调用AI生成建议
        default_type_options = [
            "技术分享",
            "产品介绍",
            "学术报告",
            "商业汇报",
            "教学课件",
            "项目展示",
            "数据分析",
            "综合介绍",
        ]

        return templates.TemplateResponse(
            "project_requirements.html",
            {
                "request": request,
                "project": project,
                "ai_suggestions": {"type_options": default_type_options},
            },
        )

    except Exception as e:
        return templates.TemplateResponse("error.html", {"request": request, "error": str(e)})


# 移除AI生成需求建议的API端点，改为使用默认选项


@router.get("/projects/{project_id}/outline-stream")
async def stream_outline_generation(
    project_id: str, user: User = Depends(get_current_user_required)
):
    """Stream outline generation for a project"""
    try:
        project = await ppt_service.project_manager.get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        if not _user_can_access_project(user, project):
            raise HTTPException(status_code=403, detail="Access denied")

        async def generate():
            try:
                async for chunk in ppt_service.generate_outline_streaming(project_id):
                    yield chunk
            except Exception as e:
                import json

                error_response = {"error": str(e)}
                yield f"data: {json.dumps(error_response)}\n\n"

        return StreamingResponse(generate(), media_type="text/plain")

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/projects/{project_id}/generate-outline")
async def generate_outline(project_id: str, user: User = Depends(get_current_user_required)):
    """Generate outline for a project (non-streaming)"""
    try:
        project = await ppt_service.project_manager.get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        if not _user_can_access_project(user, project):
            raise HTTPException(status_code=403, detail="Access denied")

        # Check if project has confirmed requirements
        if not project.confirmed_requirements:
            return {
                "status": "error",
                "error": "项目需求尚未确认，请先完成需求确认步骤",
            }

        # Create PPTGenerationRequest from project data
        confirmed_requirements = project.confirmed_requirements

        # Extract network_mode from project metadata
        network_mode = False
        if project.project_metadata and isinstance(project.project_metadata, dict):
            network_mode = project.project_metadata.get("network_mode", False)

        project_request = PPTGenerationRequest(
            scenario=project.scenario,
            topic=confirmed_requirements.get("topic", project.topic),
            requirements=project.requirements,
            language="zh",  # Default language
            network_mode=network_mode,
            target_audience=confirmed_requirements.get("target_audience", "普通大众"),
            ppt_style=confirmed_requirements.get("ppt_style", "general"),
            custom_style_prompt=confirmed_requirements.get("custom_style_prompt"),
            description=confirmed_requirements.get("description"),
        )

        # Extract page count settings from confirmed requirements
        page_count_settings = confirmed_requirements.get("page_count_settings", {})

        # Generate outline using AI with page count settings
        outline = await ppt_service.generate_outline(project_request, page_count_settings)

        # Convert outline to dict format
        outline_dict = {
            "title": outline.title,
            "slides": outline.slides,
            "metadata": outline.metadata,
        }

        # Format as JSON
        import json

        formatted_json = json.dumps(outline_dict, ensure_ascii=False, indent=2)

        # Update outline generation stage
        await ppt_service._update_outline_generation_stage(project_id, outline_dict)

        return {
            "status": "success",
            "outline_content": formatted_json,
            "message": "Outline generated successfully",
        }

    except Exception as e:
        logger.error(f"Error generating outline: {e}")
        return {"status": "error", "error": str(e)}


@router.post("/projects/{project_id}/regenerate-outline")
async def regenerate_outline(project_id: str, user: User = Depends(get_current_user_required)):
    """Regenerate outline for a project (overwrites existing outline)"""
    try:
        project = await ppt_service.project_manager.get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        if not _user_can_access_project(user, project):
            raise HTTPException(status_code=403, detail="Access denied")

        # Check if project has confirmed requirements
        if not project.confirmed_requirements:
            return {
                "status": "error",
                "error": "项目需求尚未确认，请先完成需求确认步骤",
            }

        # Create project request from confirmed requirements
        confirmed_requirements = project.confirmed_requirements
        project_request = PPTGenerationRequest(
            scenario=confirmed_requirements.get("scenario", "general"),
            topic=confirmed_requirements.get("topic", project.topic),
            requirements=confirmed_requirements.get("requirements", project.requirements),
            language="zh",  # Default language
            network_mode=confirmed_requirements.get("network_mode", False),
            target_audience=confirmed_requirements.get("target_audience", "普通大众"),
            ppt_style=confirmed_requirements.get("ppt_style", "general"),
            custom_style_prompt=confirmed_requirements.get("custom_style_prompt"),
            description=confirmed_requirements.get("description"),
        )

        # Extract page count settings from confirmed requirements
        page_count_settings = confirmed_requirements.get("page_count_settings", {})

        # Check if this is a file-based project
        is_file_project = confirmed_requirements.get("file_path") is not None

        if is_file_project:
            # Use file-based outline generation
            file_request = FileOutlineGenerationRequest(
                file_path=confirmed_requirements.get("file_path"),
                filename=confirmed_requirements.get("filename", "uploaded_file"),
                topic=project_request.topic,
                scenario=project_request.scenario,
                requirements=confirmed_requirements.get("requirements", ""),
                target_audience=confirmed_requirements.get("target_audience", "普通大众"),
                page_count_mode=page_count_settings.get("mode", "ai_decide"),
                min_pages=page_count_settings.get("min_pages", 5),
                max_pages=page_count_settings.get("max_pages", 20),
                fixed_pages=page_count_settings.get("fixed_pages", 10),
                ppt_style=confirmed_requirements.get("ppt_style", "general"),
                custom_style_prompt=confirmed_requirements.get("custom_style_prompt"),
                file_processing_mode=confirmed_requirements.get(
                    "file_processing_mode", "markitdown"
                ),
                content_analysis_depth=confirmed_requirements.get(
                    "content_analysis_depth", "standard"
                ),
            )


            result = await ppt_service.generate_outline_from_file(file_request)

            # Update outline generation stage
            await ppt_service._update_outline_generation_stage(project_id, result["outline_dict"])

            return {
                "status": "success",
                "outline_content": result["outline_content"],
                "message": "File-based outline regenerated successfully",
            }
        else:
            # Use standard outline generation
            outline = await ppt_service.generate_outline(project_request, page_count_settings)

            # Convert outline to dict format
            outline_dict = {
                "title": outline.title,
                "slides": outline.slides,
                "metadata": outline.metadata,
            }

            # Format as JSON
            import json

            formatted_json = json.dumps(outline_dict, ensure_ascii=False, indent=2)

            # Update outline generation stage
            await ppt_service._update_outline_generation_stage(project_id, outline_dict)

            return {
                "status": "success",
                "outline_content": formatted_json,
                "message": "Outline regenerated successfully",
            }

    except Exception as e:
        logger.error(f"Error regenerating outline: {e}")
        return {"status": "error", "error": str(e)}


@router.post("/projects/{project_id}/generate-file-outline")
async def generate_file_outline(project_id: str, user: User = Depends(get_current_user_required)):
    """Generate outline from uploaded file (non-streaming)"""
    try:
        project = await ppt_service.project_manager.get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        if not _user_can_access_project(user, project):
            raise HTTPException(status_code=403, detail="Access denied")

        # Check if project has file-generated outline
        file_generated_outline = None

        # 首先检查项目的outline字段
        if project.outline and project.outline.get("slides"):
            # 检查是否是从文件生成的大纲
            metadata = project.outline.get("metadata", {})
            if metadata.get("generated_with_summeryfile") or metadata.get("generated_with_file"):
                file_generated_outline = project.outline
                logger.info(
                    f"Project {project_id} has file-generated outline in project.outline, using it"
                )

        # 如果项目outline中没有，再检查confirmed_requirements
        if (
            not file_generated_outline
            and project.confirmed_requirements
            and project.confirmed_requirements.get("file_generated_outline")
        ):
            file_generated_outline = project.confirmed_requirements["file_generated_outline"]
            logger.info(
                f"Project {project_id} has file-generated outline in confirmed_requirements, using it"
            )

        # If no existing outline but file upload is configured, wait a bit and check again
        if (
            not file_generated_outline
            and project.confirmed_requirements
            and project.confirmed_requirements.get("content_source") == "file"
        ):
            logger.info(
                f"Project {project_id} has file upload but no outline yet, waiting for file processing..."
            )

            # Wait for file processing to complete (it should be done during requirements confirmation)
            import asyncio

            max_wait_time = 10  # Maximum wait time in seconds
            wait_interval = 1  # Check every 1 second

            for i in range(max_wait_time):
                await asyncio.sleep(wait_interval)

                # Refresh project data
                project = await ppt_service.project_manager.get_project(project_id)
                if project.confirmed_requirements and project.confirmed_requirements.get(
                    "file_generated_outline"
                ):
                    file_generated_outline = project.confirmed_requirements[
                        "file_generated_outline"
                    ]
                    logger.info(
                        f"Project {project_id} file outline found after waiting {i+1} seconds"
                    )
                    break

            if not file_generated_outline:
                logger.warning(
                    f"Project {project_id} file outline not found after waiting {max_wait_time} seconds"
                )

        if file_generated_outline:
            # Return the existing file-generated outline
            import json

            existing_outline = {
                "title": file_generated_outline.get("title", project.topic),
                "slides": file_generated_outline.get("slides", []),
                "metadata": file_generated_outline.get("metadata", {}),
            }

            # Ensure metadata includes correct identification
            if "metadata" not in existing_outline:
                existing_outline["metadata"] = {}
            existing_outline["metadata"]["generated_with_summeryfile"] = True
            existing_outline["metadata"]["generated_at"] = time.time()

            formatted_json = json.dumps(existing_outline, ensure_ascii=False, indent=2)

            # Update outline generation stage
            await ppt_service._update_outline_generation_stage(project_id, existing_outline)

            return {
                "status": "success",
                "outline_content": formatted_json,
                "message": "File outline generated successfully",
            }
        else:
            # Check if there's an uploaded file that needs processing
            if project.confirmed_requirements and (
                project.confirmed_requirements.get("uploaded_files")
                or project.confirmed_requirements.get("content_source") == "file"
            ):
                logger.info(
                    f"Project {project_id} has uploaded files, starting file outline generation"
                )

                # Start file outline generation using summeryfile
                try:
                    # Create a request object for file outline generation
                    from ..api.models import FileOutlineGenerationRequest

                    # Get file information from confirmed requirements
                    uploaded_files = project.confirmed_requirements.get("uploaded_files", [])
                    if uploaded_files:
                        file_info = uploaded_files[0]  # Use first file
                        # 使用确认的要求或项目创建时的要求作为fallback
                        confirmed_reqs = project.confirmed_requirements.get("requirements", "")
                        project_reqs = project.requirements or ""
                        final_reqs = confirmed_reqs or project_reqs

                        file_request = FileOutlineGenerationRequest(
                            filename=file_info.get("filename", "uploaded_file"),
                            file_path=file_info.get("file_path", ""),
                            topic=project.topic,
                            scenario="general",
                            requirements=final_reqs,
                            target_audience=project.confirmed_requirements.get(
                                "target_audience", "普通大众"
                            ),
                            page_count_mode=project.confirmed_requirements.get(
                                "page_count_settings", {}
                            ).get("mode", "ai_decide"),
                            min_pages=project.confirmed_requirements.get(
                                "page_count_settings", {}
                            ).get("min_pages", 8),
                            max_pages=project.confirmed_requirements.get(
                                "page_count_settings", {}
                            ).get("max_pages", 15),
                            fixed_pages=project.confirmed_requirements.get(
                                "page_count_settings", {}
                            ).get("fixed_pages", 10),
                            ppt_style=project.confirmed_requirements.get("ppt_style", "general"),
                            custom_style_prompt=project.confirmed_requirements.get(
                                "custom_style_prompt"
                            ),
                            file_processing_mode=project.confirmed_requirements.get(
                                "file_processing_mode", "markitdown"
                            ),
                            content_analysis_depth=project.confirmed_requirements.get(
                                "content_analysis_depth", "standard"
                            ),
                        )

                        # Generate outline from file using summeryfile
                        outline_response = await ppt_service.generate_outline_from_file(
                            file_request
                        )

                        if outline_response.success and outline_response.outline:
                            # Format the generated outline
                            import json

                            formatted_outline = outline_response.outline

                            # Ensure metadata includes correct identification
                            if "metadata" not in formatted_outline:
                                formatted_outline["metadata"] = {}
                            formatted_outline["metadata"]["generated_with_summeryfile"] = True
                            formatted_outline["metadata"]["generated_at"] = time.time()

                            formatted_json = json.dumps(
                                formatted_outline, ensure_ascii=False, indent=2
                            )

                            # Update outline generation stage
                            await ppt_service._update_outline_generation_stage(
                                project_id, formatted_outline
                            )

                            return {
                                "status": "success",
                                "outline_content": formatted_json,
                                "message": "File outline generated successfully",
                            }
                        else:
                            error_msg = (
                                outline_response.error
                                if hasattr(outline_response, "error")
                                else "Unknown error"
                            )
                            return {
                                "status": "error",
                                "error": f"Failed to generate outline from uploaded file: {error_msg}",
                            }
                    else:
                        return {
                            "status": "error",
                            "error": "No uploaded file information found in project requirements.",
                        }

                except Exception as gen_error:
                    logger.error(f"Error generating outline from file: {gen_error}")
                    return {
                        "status": "error",
                        "error": f"Failed to generate outline from file: {str(gen_error)}",
                    }
            else:
                # No file outline found and no uploaded files
                return {
                    "status": "error",
                    "error": "No file outline found. Please ensure you uploaded a file during requirements confirmation.",
                }

    except Exception as e:
        logger.error(f"Error generating file outline: {e}")
        return {"status": "error", "error": str(e)}


@router.post("/projects/{project_id}/update-outline")
async def update_project_outline(
    project_id: str, request: Request, user: User = Depends(get_current_user_required)
):
    """Update project outline content"""
    try:
        data = await request.json()
        outline_content = data.get("outline_content", "")

        # Access check
        project = await ppt_service.project_manager.get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        if not _user_can_access_project(user, project):
            raise HTTPException(status_code=403, detail="Access denied")

        success = await ppt_service.update_project_outline(project_id, outline_content)
        if success:
            return {"status": "success", "message": "Outline updated"}
        else:
            raise HTTPException(status_code=500, detail="Failed to update outline")

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/projects/{project_id}/confirm-outline")
async def confirm_project_outline(project_id: str, user: User = Depends(get_current_user_required)):
    """Confirm project outline and enable PPT generation"""
    try:
        # Access check
        project = await ppt_service.project_manager.get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        if not _user_can_access_project(user, project):
            raise HTTPException(status_code=403, detail="Access denied")

        success = await ppt_service.confirm_project_outline(project_id)
        if success:
            return {"status": "success", "message": "Outline confirmed"}
        else:
            raise HTTPException(status_code=500, detail="Failed to confirm outline")

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/projects/{project_id}/todo-editor")
async def web_project_todo_editor(
    request: Request,
    project_id: str,
    auto_start: bool = False,
    user: User = Depends(get_current_user_required),
):
    """Project TODO board with editor"""
    try:
        project = await ppt_service.project_manager.get_project(project_id)
        if not project:
            return templates.TemplateResponse(
                "error.html", {"request": request, "error": "Project not found"}
            )
        if not _user_can_access_project(user, project):
            raise HTTPException(status_code=403, detail="Access denied")

        return templates.TemplateResponse(
            "todo_board_with_editor.html",
            {
                "request": request,
                "todo_board": project.todo_board,
                "project": project,
                "auto_start": auto_start,
            },
        )

    except Exception as e:
        return templates.TemplateResponse("error.html", {"request": request, "error": str(e)})


@router.post("/projects/{project_id}/confirm-requirements")
async def confirm_project_requirements(
    request: Request,
    project_id: str,
    topic: str = Form(...),
    audience_type: str = Form(...),
    custom_audience: str = Form(None),
    page_count_mode: str = Form("ai_decide"),
    min_pages: int = Form(8),
    max_pages: int = Form(15),
    fixed_pages: int = Form(10),
    ppt_style: str = Form("general"),
    custom_style_prompt: str = Form(None),
    description: str = Form(None),
    content_source: str = Form("manual"),
    file_upload: Optional[UploadFile] = File(None),
    file_processing_mode: str = Form("markitdown"),
    content_analysis_depth: str = Form("standard"),
    user: User = Depends(get_current_user_required),
):
    """Confirm project requirements and generate TODO list"""
    try:
        # Get project to access original requirements
        project = await ppt_service.project_manager.get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        if not _user_can_access_project(user, project):
            raise HTTPException(status_code=403, detail="Access denied")

        # Process audience information
        target_audience = audience_type
        if audience_type == "自定义" and custom_audience:
            target_audience = custom_audience

        # Handle file upload if content source is file
        file_outline = None
        if content_source == "file" and file_upload:
            # Process uploaded file and generate outline
            # 使用项目创建时的具体要求而不是description
            file_outline = await _process_uploaded_file_for_outline(
                file_upload,
                topic,
                target_audience,
                page_count_mode,
                min_pages,
                max_pages,
                fixed_pages,
                ppt_style,
                custom_style_prompt,
                file_processing_mode,
                content_analysis_depth,
                project.requirements,
            )

            # Update topic if it was extracted from file
            if file_outline and file_outline.get("title") and not topic.strip():
                topic = file_outline["title"]

        # Process page count settings
        page_count_settings = {
            "mode": page_count_mode,
            "min_pages": min_pages if page_count_mode == "custom_range" else None,
            "max_pages": max_pages if page_count_mode == "custom_range" else None,
            "fixed_pages": fixed_pages if page_count_mode == "fixed" else None,
        }

        # Update project with confirmed requirements
        confirmed_requirements = {
            "topic": topic,
            "requirements": project.requirements,  # 使用项目创建时的具体要求
            "target_audience": target_audience,
            "audience_type": audience_type,
            "custom_audience": custom_audience if audience_type == "自定义" else None,
            "page_count_settings": page_count_settings,
            "ppt_style": ppt_style,
            "custom_style_prompt": (custom_style_prompt if ppt_style == "custom" else None),
            "description": description,
            "content_source": content_source,
            "file_processing_mode": (file_processing_mode if content_source == "file" else None),
            "content_analysis_depth": (
                content_analysis_depth if content_source == "file" else None
            ),
            "file_generated_outline": file_outline,
        }

        # Store confirmed requirements in project
        # 直接确认需求并更新TODO板，无需AI生成待办清单
        success = await ppt_service.confirm_requirements_and_update_workflow(
            project_id, confirmed_requirements
        )

        if not success:
            raise Exception("需求确认失败")

        # Return JSON success response for AJAX request
        from fastapi.responses import JSONResponse

        return JSONResponse(
            {
                "status": "success",
                "message": "需求确认完成",
                "redirect_url": f"/projects/{project_id}/todo",
            }
        )

    except Exception as e:
        from fastapi.responses import JSONResponse

        return JSONResponse({"status": "error", "message": str(e)}, status_code=500)


@router.get("/projects/{project_id}/stage-stream/{stage_id}")
async def stream_stage_response(
    project_id: str, stage_id: str, user: User = Depends(get_current_user_required)
):
    """Stream AI response for a complete stage"""

    async def generate_stage_stream():
        try:
            # Get project and stage info
            project = await ppt_service.project_manager.get_project(project_id)
            if not project:
                yield f"data: {json.dumps({'error': 'Project not found'})}\n\n"
                return

            if not project.confirmed_requirements:
                yield f"data: {json.dumps({'error': 'Project requirements not confirmed'})}\n\n"
                return

            todo_board = await ppt_service.get_project_todo_board(project_id)
            if not todo_board:
                yield f"data: {json.dumps({'error': 'TODO board not found'})}\n\n"
                return

            # Find the stage
            stage = None
            for s in todo_board.stages:
                if s.id == stage_id:
                    stage = s
                    break

            if not stage:
                yield f"data: {json.dumps({'error': 'Stage not found'})}\n\n"
                return

            # Extract confirmed requirements from project
            confirmed_requirements = project.confirmed_requirements

            # Check if stage is already running or completed
            if stage.status == "running":
                yield f"data: {json.dumps({'error': 'Stage is already running'})}\n\n"
                return
            elif stage.status == "completed":
                yield f"data: {json.dumps({'error': 'Stage is already completed'})}\n\n"
                return

            # Update stage status to running
            await ppt_service.project_manager.update_stage_status(
                project_id, stage_id, "running", 0.0
            )

            # Execute the complete stage using the enhanced service
            try:
                if stage_id == "outline_generation":
                    response_content = await ppt_service._execute_outline_generation(
                        project_id,
                        confirmed_requirements,
                        ppt_service._load_prompts_md_system_prompt(),
                    )
                elif stage_id == "ppt_creation":
                    response_content = await ppt_service._execute_ppt_creation(
                        project_id,
                        confirmed_requirements,
                        ppt_service._load_prompts_md_system_prompt(),
                    )
                else:
                    # Fallback for other stages
                    response_content = await ppt_service._execute_general_stage(
                        project_id, stage_id, confirmed_requirements
                    )

                # Stream the response word by word for better UX
                if isinstance(response_content, dict):
                    content_text = response_content.get("message", str(response_content))
                else:
                    content_text = str(response_content)

                words = content_text.split()
                for i, word in enumerate(words):
                    yield f"data: {json.dumps({'content': word + ' ', 'done': False})}\n\n"
                    await asyncio.sleep(0.05)  # Small delay for streaming effect

            except Exception:
                # Fallback to basic stage execution
                prompt = f"""
作为PPT生成助手，请完成以下阶段任务：

项目主题：{project.topic}
项目场景：{project.scenario}
项目要求：{project.requirements or '无特殊要求'}

当前阶段：{stage.name}
阶段描述：{stage.description}

请根据以上信息完成当前阶段的完整任务，并提供详细的执行结果。
"""

                # Stream AI response using real streaming
                async for chunk in ppt_service.ai_provider.stream_text_completion(
                    prompt=prompt, max_tokens=2000, temperature=0.7
                ):
                    if chunk:
                        yield f"data: {json.dumps({'content': chunk, 'done': False})}\n\n"

            # Update stage status to completed
            await ppt_service.project_manager.update_stage_status(
                project_id, stage_id, "completed", 100.0
            )

            # Send completion signal
            yield f"data: {json.dumps({'content': '', 'done': True})}\n\n"

        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    return StreamingResponse(
        generate_stage_stream(),
        media_type="text/plain",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive"},
    )


@router.get("/projects/{project_id}/edit", response_class=HTMLResponse)
async def edit_project_ppt(
    request: Request, project_id: str, user: User = Depends(get_current_user_required)
):
    """Edit PPT slides with advanced editor"""
    try:
        project = await ppt_service.project_manager.get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")

        # 允许编辑器在Slide生成过程中显示，提供更好的用户体验
        # 如果没有slides_data，创建一个空的结构供编辑器使用
        if not project.slides_data:
            project.slides_data = []

        return templates.TemplateResponse(
            "project_slides_editor.html",
            {
                "request": request,
                "project": project,
                "aspect_ratio_settings": get_aspect_ratio_settings(),
            },
        )

    except Exception as e:
        return templates.TemplateResponse("error.html", {"request": request, "error": str(e)})


@router.post("/api/projects/{project_id}/update-html")
async def update_project_html(
    project_id: str, request: Request, user: User = Depends(get_current_user_required)
):
    """Update project HTML content and mark all slides as user-edited"""
    try:
        data = await request.json()
        slides_html = data.get("slides_html", "")

        project = await ppt_service.project_manager.get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        if not _user_can_access_project(user, project):
            raise HTTPException(status_code=403, detail="Access denied")

        # Update project HTML
        project.slides_html = slides_html
        project.updated_at = time.time()

        # 解析HTML内容，提取各个页面并标记为用户编辑
        if project.slides_data and slides_html:
            try:
                # 解析HTML内容，提取各个页面
                updated_slides_data = await _extract_slides_from_html(
                    slides_html, project.slides_data
                )

                # 标记所有页面为用户编辑状态
                for slide_data in updated_slides_data:
                    slide_data["is_user_edited"] = True

                # 更新项目的slides_data
                project.slides_data = updated_slides_data

                logger.info(
                    f"Marked {len(updated_slides_data)} slides as user-edited for project {project_id}"
                )

            except Exception as parse_error:
                logger.warning(f"Failed to parse HTML content for slide extraction: {parse_error}")
                # 如果解析失败，至少标记现有的slides_data为用户编辑
                if project.slides_data:
                    for slide_data in project.slides_data:
                        slide_data["is_user_edited"] = True

        # 保存更新的HTML和slides_data到数据库
        try:
            from ..services.db_project_manager import DatabaseProjectManager

            db_manager = DatabaseProjectManager()

            # 保存幻灯片HTML和数据到数据库
            save_success = await db_manager.save_project_slides(
                project_id, project.slides_html, project.slides_data or []
            )

            if save_success:
                logger.info(
                    f"Successfully saved updated HTML and slides data to database for project {project_id}"
                )
            else:
                logger.error(
                    f"Failed to save updated HTML and slides data to database for project {project_id}"
                )

        except Exception as save_error:
            logger.error(
                f"Exception while saving updated HTML and slides data to database: {save_error}"
            )
            # 继续返回成功，因为内存中的数据已经更新

        return {
            "status": "success",
            "message": "HTML updated successfully and slides marked as user-edited",
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/api/projects/{project_id}")
async def get_project_data(project_id: str, user: User = Depends(get_current_user_required)):
    """Get project data for real-time updates"""
    try:
        project = await ppt_service.project_manager.get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        if not _user_can_access_project(user, project):
            raise HTTPException(status_code=403, detail="Access denied")

        return {
            "project_id": project.project_id,
            "title": project.title,
            "status": project.status,
            "slides_data": project.slides_data or [],
            "slides_count": len(project.slides_data) if project.slides_data else 0,
            "updated_at": project.updated_at,
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/api/projects/{project_id}/slides")
async def update_project_slides(
    project_id: str, request: Request, user: User = Depends(get_current_user_required)
):
    """Update project slides data"""
    try:
        logger.info(f"🔄 开始更新项目 {project_id} 的幻灯片数据")

        data = await request.json()
        slides_data = data.get("slides_data", [])

        logger.info(f"📊 接收到 {len(slides_data)} 页幻灯片数据")

        project = await ppt_service.project_manager.get_project(project_id)
        if not project:
            logger.error(f"❌ 项目 {project_id} 不存在")
            raise HTTPException(status_code=404, detail="Project not found")
        if not _user_can_access_project(user, project):
            raise HTTPException(status_code=403, detail="Access denied")

        logger.info("📝 更新项目幻灯片数据...")

        # Update project slides data
        project.slides_data = slides_data
        project.updated_at = time.time()

        # Regenerate combined HTML
        if slides_data:
            # 安全地获取大纲标题
            outline_title = project.title
            if project.outline:
                if isinstance(project.outline, dict):
                    outline_title = project.outline.get("title", project.title)
                elif hasattr(project.outline, "title"):
                    outline_title = project.outline.title

            project.slides_html = ppt_service._combine_slides_to_full_html(
                slides_data, outline_title
            )

        # 标记所有幻灯片为用户编辑状态
        for i, slide_data in enumerate(project.slides_data):
            slide_data["is_user_edited"] = True

        # 保存更新的幻灯片数据到数据库
        save_success = False
        save_error_message = None

        try:
            from ..services.db_project_manager import DatabaseProjectManager

            db_manager = DatabaseProjectManager()

            # 保存幻灯片数据到数据库
            save_success = await db_manager.save_project_slides(
                project_id, project.slides_html or "", project.slides_data
            )

            if save_success:
                logger.info(
                    f"Successfully saved updated slides data to database for project {project_id}"
                )
            else:
                logger.error(
                    f"Failed to save updated slides data to database for project {project_id}"
                )
                save_error_message = "Failed to save slides data to database"

        except Exception as save_error:
            logger.error(f"❌ 保存幻灯片数据到数据库时发生异常: {save_error}")
            import traceback

            traceback.print_exc()
            save_success = False
            save_error_message = str(save_error)

        # 根据保存结果返回相应的响应
        if save_success:
            return {
                "status": "success",
                "success": True,
                "message": "Slides updated and saved to database successfully",
            }
        else:
            # 即使数据库保存失败，内存中的数据已经更新，所以仍然返回成功，但包含警告信息
            return {
                "status": "success",
                "success": True,
                "message": "Slides updated in memory successfully",
                "warning": f"Database save failed: {save_error_message}",
                "database_saved": False,
            }

    except Exception as e:
        logger.error(f"❌ 更新项目幻灯片数据时发生错误: {e}")
        import traceback

        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/api/projects/{project_id}/regenerate-html")
async def regenerate_project_html(project_id: str, request: Request = None):
    """Regenerate project HTML with fixed encoding"""
    try:
        project = await ppt_service.project_manager.get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        # Access control: only owner or admin can regenerate
        user = (
            getattr(request.state, "user", None)
            if isinstance(locals().get("request"), Request)
            else None
        )
        if user and not _user_can_access_project(user, project):
            raise HTTPException(status_code=403, detail="Access denied")

        if not project.slides_data:
            raise HTTPException(status_code=400, detail="No slides data found")

        # Regenerate combined HTML using the fixed method
        # 安全地获取大纲标题
        outline_title = project.title
        if project.outline:
            if isinstance(project.outline, dict):
                outline_title = project.outline.get("title", project.title)
            elif hasattr(project.outline, "title"):
                outline_title = project.outline.title

        project.slides_html = ppt_service._combine_slides_to_full_html(
            project.slides_data, outline_title
        )

        project.updated_at = time.time()

        # 保存重新生成的HTML到数据库
        try:
            from ..services.db_project_manager import DatabaseProjectManager

            db_manager = DatabaseProjectManager()

            # 保存幻灯片数据到数据库
            save_success = await db_manager.save_project_slides(
                project_id, project.slides_html, project.slides_data
            )

            if save_success:
                logger.info(
                    f"Successfully saved regenerated HTML to database for project {project_id}"
                )
            else:
                logger.error(
                    f"Failed to save regenerated HTML to database for project {project_id}"
                )

        except Exception as save_error:
            logger.error(f"Exception while saving regenerated HTML to database: {save_error}")
            # 继续返回成功，因为内存中的数据已经更新

        return {"success": True, "message": "Project HTML regenerated successfully"}

    except Exception as e:
        return {"success": False, "error": str(e)}


@router.post("/api/projects/{project_id}/slides/{slide_number}/regenerate")
async def regenerate_slide(project_id: str, slide_number: int, request: Request = None):
    """Regenerate a specific slide"""
    try:
        project = await ppt_service.project_manager.get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        # Access control: only owner or admin can regenerate slides
        user = (
            getattr(request.state, "user", None)
            if isinstance(locals().get("request"), Request)
            else None
        )
        if user and not _user_can_access_project(user, project):
            raise HTTPException(status_code=403, detail="Access denied")

        if not project.outline:
            raise HTTPException(status_code=400, detail="Project outline not found")

        if not project.confirmed_requirements:
            raise HTTPException(status_code=400, detail="Project requirements not confirmed")

        # Handle different outline structures
        if isinstance(project.outline, dict):
            slides = project.outline.get("slides", [])
        else:
            # If outline is a PPTOutline object
            slides = project.outline.slides if hasattr(project.outline, "slides") else []

        if slide_number < 1 or slide_number > len(slides):
            raise HTTPException(status_code=400, detail="Invalid slide number")

        slide_data = slides[slide_number - 1]

        # Load system prompt
        system_prompt = ppt_service._load_prompts_md_system_prompt()

        # Ensure project has a global template selected (use default if none selected)
        selected_template = await ppt_service._ensure_global_master_template_selected(project_id)

        # Regenerate the slide using template-based generation if template is available
        if selected_template:
            logger.info(
                f"Regenerating slide {slide_number} using template: {selected_template['template_name']}"
            )
            new_html_content = await ppt_service._generate_slide_with_template(
                slide_data,
                selected_template,
                slide_number,
                len(slides),
                project.confirmed_requirements,
            )
        else:
            # Fallback to original generation method if no template available
            logger.warning(
                f"No template available for project {project_id}, using fallback generation"
            )
            new_html_content = await ppt_service._generate_single_slide_html_with_prompts(
                slide_data,
                project.confirmed_requirements,
                system_prompt,
                slide_number,
                len(slides),
                slides,
                project.slides_data,
                project_id=project_id,
            )

        # Update the slide in project data
        if not project.slides_data:
            project.slides_data = []

        # Ensure slides_data has enough entries
        while len(project.slides_data) < slide_number:
            new_page_number = len(project.slides_data) + 1
            project.slides_data.append(
                {
                    "page_number": new_page_number,
                    "title": f"第{new_page_number}页",
                    "html_content": "<div>待生成</div>",
                    "slide_type": "content",
                    "content_points": [],
                    "is_user_edited": False,
                }
            )

        # Update the specific slide - 保持与现有数据结构一致
        existing_slide = (
            project.slides_data[slide_number - 1]
            if slide_number <= len(project.slides_data)
            else {}
        )

        # 更新幻灯片数据，保留现有字段并确保必要字段存在
        updated_slide = {
            "page_number": slide_number,
            "title": slide_data.get("title", f"第{slide_number}页"),
            "html_content": new_html_content,
            "slide_type": slide_data.get("slide_type", existing_slide.get("slide_type", "content")),
            "content_points": slide_data.get(
                "content_points", existing_slide.get("content_points", [])
            ),
            "is_user_edited": existing_slide.get("is_user_edited", False),
            # 保留其他可能存在的字段
            **{
                k: v
                for k, v in existing_slide.items()
                if k
                not in [
                    "page_number",
                    "title",
                    "html_content",
                    "slide_type",
                    "content_points",
                    "is_user_edited",
                ]
            },
        }

        project.slides_data[slide_number - 1] = updated_slide

        # Regenerate combined HTML
        outline_title = project.title
        if isinstance(project.outline, dict):
            outline_title = project.outline.get("title", project.title)
        elif hasattr(project.outline, "title"):
            outline_title = project.outline.title

        project.slides_html = ppt_service._combine_slides_to_full_html(
            project.slides_data, outline_title
        )

        project.updated_at = time.time()

        # 保存更新后的幻灯片数据到数据库
        try:
            from ..services.db_project_manager import DatabaseProjectManager

            db_manager = DatabaseProjectManager()

            # 只保存单个重新生成的幻灯片，而不是整个项目的幻灯片数据
            # 这样可以避免删除所有幻灯片再重新创建的问题
            save_success = await db_manager.save_single_slide(
                project_id, slide_number - 1, updated_slide  # 转换为0基索引
            )

            if save_success:
                logger.info(
                    f"Successfully saved regenerated slide {slide_number} to database for project {project_id}"
                )

                # 同时更新项目的slides_html字段
                await db_manager.update_project_data(
                    project_id,
                    {
                        "slides_html": project.slides_html,
                        "updated_at": project.updated_at,
                    },
                )
            else:
                logger.error(
                    f"Failed to save regenerated slide {slide_number} to database for project {project_id}"
                )

        except Exception as save_error:
            logger.error(f"Exception while saving regenerated slide to database: {save_error}")
            # 继续返回成功，因为内存中的数据已经更新

        return {
            "success": True,
            "message": f"Slide {slide_number} regenerated successfully",
            "slide_data": project.slides_data[slide_number - 1],
        }

    except Exception as e:
        return {"success": False, "error": str(e)}


@router.post("/api/ai/slide-edit")
async def ai_slide_edit(
    request: AISlideEditRequest, user: User = Depends(get_current_user_required)
):
    """AI编辑幻灯片接口"""
    try:
        # 获取AI提供者
        ai_provider = get_ai_provider()

        # 构建AI编辑上下文
        outline_info = ""
        if request.slideOutline:
            outline_info = f"""
当前幻灯片大纲信息：
- 幻灯片类型：{request.slideOutline.get('slide_type', '未知')}
- 描述：{request.slideOutline.get('description', '无')}
- 要点：{', '.join(request.slideOutline.get('content_points', [])) if request.slideOutline.get('content_points') else '无'}
"""

        context = f"""
你是一位专业的Slide设计师和编辑助手。用户想要对当前幻灯片进行编辑修改。

当前幻灯片信息：
- 页码：第{request.slideIndex}页
- 标题：{request.slideTitle}
- 项目主题：{request.projectInfo.get('title', '未知')}
- 项目场景：{request.projectInfo.get('scenario', '未知')}
{outline_info}
用户的编辑要求：
{request.userRequest}

当前幻灯片的HTML内容：
{request.slideContent}

请根据用户的要求和幻灯片大纲信息，提供以下内容：
1. 对用户要求的理解和分析
2. 具体的修改建议
3. 如果需要，提供修改后的完整HTML代码

注意事项：
- 确保修改后的内容符合Slide演示的专业标准和大纲要求
- 生成的HTML应该是完整的，包含必要的CSS样式
- 保持1280x720的Slide标准尺寸
- 参考大纲信息中的要点和描述来优化内容
"""

        # 构建AI消息，包含对话历史
        messages = [
            AIMessage(
                role=MessageRole.SYSTEM,
                content="你是一位专业的Slide设计师和编辑助手，擅长根据用户需求修改和优化Slide内容。",
            )
        ]

        # 添加对话历史
        if request.chatHistory:
            logger.debug(f"AI编辑接收到对话历史，共 {len(request.chatHistory)} 条消息")
            for i, chat_msg in enumerate(request.chatHistory):
                role = MessageRole.USER if chat_msg.get("role") == "user" else MessageRole.ASSISTANT
                content = chat_msg.get("content", "")
                logger.debug(f"对话历史 {i+1}: {role.value} - {content[:100]}...")
                messages.append(AIMessage(role=role, content=content))
        else:
            logger.debug("AI编辑未接收到对话历史")

        # 添加当前用户请求
        messages.append(AIMessage(role=MessageRole.USER, content=context))

        # 调用AI生成回复
        response = await ai_provider.chat_completion(
            messages=messages, max_tokens=ai_config.max_tokens, temperature=0.7
        )

        ai_response = response.content

        # 检查是否包含HTML代码
        new_html_content = None
        if "```html" in ai_response:
            import re

            html_match = re.search(r"```html\s*(.*?)\s*```", ai_response, re.DOTALL)
            if html_match:
                new_html_content = html_match.group(1).strip()

        return {
            "success": True,
            "response": ai_response,
            "newHtmlContent": new_html_content,
        }

    except Exception as e:
        logger.error(f"AI编辑请求失败: {e}")
        return {
            "success": False,
            "error": str(e),
            "response": "抱歉，AI编辑服务暂时不可用。请稍后重试。",
        }


@router.post("/api/ai/slide-edit/stream")
async def ai_slide_edit_stream(
    request: AISlideEditRequest, user: User = Depends(get_current_user_required)
):
    """AI编辑幻灯片流式接口"""
    try:
        # 获取AI提供者
        ai_provider = get_ai_provider()

        # 构建AI编辑上下文
        outline_info = ""
        if request.slideOutline:
            outline_info = """
当前幻灯片大纲信息：
""" + str(request.slideOutline)

        # 构建图片信息
        images_info = ""
        if request.images and len(request.images) > 0:
            images_info = """

用户上传的图片信息：
"""
            for i, image in enumerate(request.images, 1):
                images_info += f"""
- 图片{i}：{image.get('name', '未知')}
  - URL：{image.get('url', '')}
  - 大小：{image.get('size', '未知')}
  - 说明：请分析这张图片的内容，理解用户的意图，并根据编辑要求进行相应的处理
"""

        context = f"""
你是一位专业的Slide设计师和编辑助手。用户想要对当前幻灯片进行编辑修改。

当前幻灯片信息：
- 页码：第{request.slideIndex}页
- 标题：{request.slideTitle}
- 项目主题：{request.projectInfo.get('title', '未知')}
- 项目场景：{request.projectInfo.get('scenario', '未知')}
{outline_info}{images_info}
用户的编辑要求：
{request.userRequest}

当前幻灯片的HTML内容：
{request.slideContent}

请根据用户的要求和幻灯片大纲信息，提供以下内容：
1. 对用户要求的理解和分析
2. 具体的修改建议
3. 默认提供修改后的完整HTML代码

注意事项：
- 保持原有的设计风格和布局结构
- 确保修改后的内容符合Slide演示的专业标准和大纲要求
- 如果用户要求不明确，请提供多个可选方案
- 生成的HTML应该是完整的，包含必要的CSS样式
- 保持1280x720的Slide标准尺寸
- 参考大纲信息中的要点和描述来优化内容
"""

        # 构建AI消息，包含对话历史
        messages = [
            AIMessage(
                role=MessageRole.SYSTEM,
                content="你是一位专业的Slide设计师和编辑助手，擅长根据用户需求修改和优化Slide内容。",
            )
        ]

        # 添加对话历史
        if request.chatHistory:
            logger.info(f"AI流式编辑接收到对话历史，共 {len(request.chatHistory)} 条消息")
            for i, chat_msg in enumerate(request.chatHistory):
                role = MessageRole.USER if chat_msg.get("role") == "user" else MessageRole.ASSISTANT
                content = chat_msg.get("content", "")
                logger.info(f"对话历史 {i+1}: {role.value} - {content[:100]}...")
                messages.append(AIMessage(role=role, content=content))
        else:
            logger.info("AI流式编辑未接收到对话历史")

        # 添加当前用户请求
        messages.append(AIMessage(role=MessageRole.USER, content=context))

        async def generate_ai_stream():
            try:
                # 发送开始信号
                yield f"data: {json.dumps({'type': 'start', 'content': ''})}\n\n"

                # 流式生成AI回复
                full_response = ""
                async for chunk in ai_provider.stream_chat_completion(
                    messages=messages, max_tokens=ai_config.max_tokens, temperature=0.7
                ):
                    if chunk:
                        full_response += chunk
                        yield f"data: {json.dumps({'type': 'content', 'content': chunk})}\n\n"

                # 检查是否包含HTML代码
                new_html_content = None
                if "```html" in full_response:
                    import re

                    html_match = re.search(r"```html\s*(.*?)\s*```", full_response, re.DOTALL)
                    if html_match:
                        new_html_content = html_match.group(1).strip()

                # 发送完成信号
                yield f"data: {json.dumps({'type': 'complete', 'content': '', 'newHtmlContent': new_html_content, 'fullResponse': full_response})}\n\n"

            except Exception as e:
                logger.error(f"AI流式编辑请求失败: {e}")
                yield f"data: {json.dumps({'type': 'error', 'content': '', 'error': str(e)})}\n\n"

        return StreamingResponse(
            generate_ai_stream(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Headers": "Cache-Control",
            },
        )

    except Exception as e:
        logger.error(f"AI流式编辑请求失败: {e}")
        return {
            "success": False,
            "error": str(e),
            "response": "抱歉，AI编辑服务暂时不可用。请稍后重试。",
        }


@router.post("/api/ai/regenerate-image")
async def ai_regenerate_image(
    request: AIImageRegenerateRequest, user: User = Depends(get_current_user_required)
):
    """AI重新生成图像接口 - 完全遵循enhanced_ppt_service.py的标准流程"""
    try:
        # 获取图像服务和AI提供者
        from ..services.image.image_service import get_image_service

        image_service = get_image_service()
        if not image_service:
            return {"success": False, "message": "图像服务不可用"}

        ai_provider = get_ai_provider()
        if not ai_provider:
            return {"success": False, "message": "AI提供者不可用"}

        # 获取图像配置
        from ..services.config_service import config_service

        image_config = config_service.get_config_by_category("image_service")

        # 检查是否启用图片生成服务
        enable_image_service = image_config.get("enable_image_service", False)
        if not enable_image_service:
            return {"success": False, "message": "图片生成服务未启用，请在配置中启用"}

        # 第一步：检查启用的图片来源（完全遵循PPTImageProcessor的逻辑）
        from ..services.models.slide_image_info import ImageSource

        enabled_sources = []
        if image_config.get("enable_local_images", True):
            enabled_sources.append(ImageSource.LOCAL)
        if image_config.get("enable_network_search", False):
            enabled_sources.append(ImageSource.NETWORK)
        if image_config.get("enable_ai_generation", False):
            enabled_sources.append(ImageSource.AI_GENERATED)

        if not enabled_sources:
            return {
                "success": False,
                "message": "没有启用任何图片来源，请在配置中启用至少一种图片来源",
            }

        # 初始化Slide图像处理器
        from ..services.ppt_image_processor import PPTImageProcessor

        image_processor = PPTImageProcessor(image_service=image_service, ai_provider=ai_provider)

        # 提取图像信息和幻灯片内容
        image_info = request.image_info
        slide_content = request.slide_content

        # 构建幻灯片数据结构（仅用于调试以避免引入未使用局部变量）
        if logger.isEnabledFor(logging.DEBUG):
            slide_data = {
                "title": slide_content.get("title", ""),
                "content_points": [slide_content.get("title", "")],  # 简化的内容点
            }

            # 构建确认需求结构（仅调试）
            confirmed_requirements = {
                "project_topic": request.project_topic,
                "project_scenario": request.project_scenario,
            }
            # Use debug variables to help debugging and satisfy linters
            logger.debug(f"Debug slide_data: {slide_data}")
            logger.debug(f"Debug confirmed_requirements: {confirmed_requirements}")

        # 第二步：直接创建图像重新生成需求（跳过AI配图适用性判断）
        logger.info(f"开始图片重新生成，启用的来源: {[source.value for source in enabled_sources]}")

        # 分析原图像的用途和上下文
        image_context = await analyze_image_context(
            image_info, slide_content, request.project_topic, request.project_scenario
        )

        # 根据启用的来源和配置，智能选择最佳的图片来源
        selected_source = select_best_image_source(enabled_sources, image_config, image_context)

        # 创建图像需求对象（直接生成，不需要AI判断是否适合配图）
        from ..services.models.slide_image_info import ImagePurpose, ImageRequirement

        # 将字符串用途转换为ImagePurpose枚举
        purpose_str = image_context.get("image_purpose", "illustration")
        purpose_mapping = {
            "background": ImagePurpose.BACKGROUND,
            "icon": ImagePurpose.ICON,
            "chart_support": ImagePurpose.CHART_SUPPORT,
            "decoration": ImagePurpose.DECORATION,
            "illustration": ImagePurpose.ILLUSTRATION,
        }
        purpose = purpose_mapping.get(purpose_str, ImagePurpose.ILLUSTRATION)

        requirement = ImageRequirement(
            source=selected_source,
            count=1,
            purpose=purpose,
            description=f"重新生成图像: {image_info.get('alt', '')} - {request.project_topic}",
            priority=5,  # 高优先级，因为是用户明确请求的重新生成
        )

        logger.info(f"选择图片来源: {selected_source.value}, 用途: {purpose.value}")

        # 第三步：直接处理图片生成（单个需求）
        from ..services.models.slide_image_info import SlideImagesCollection

        images_collection = SlideImagesCollection(page_number=request.slide_index + 1, images=[])

        # 根据选择的来源处理图片生成
        if requirement.source == ImageSource.LOCAL and ImageSource.LOCAL in enabled_sources:
            local_images = await image_processor._process_local_images(
                requirement,
                request.project_topic,
                request.project_scenario,
                slide_content.get("title", ""),
                slide_content.get("title", ""),
            )
            images_collection.images.extend(local_images)

        elif requirement.source == ImageSource.NETWORK and ImageSource.NETWORK in enabled_sources:
            network_images = await image_processor._process_network_images(
                requirement,
                request.project_topic,
                request.project_scenario,
                slide_content.get("title", ""),
                slide_content.get("title", ""),
                image_config,
            )
            images_collection.images.extend(network_images)

        elif (
            requirement.source == ImageSource.AI_GENERATED
            and ImageSource.AI_GENERATED in enabled_sources
        ):
            ai_images = await image_processor._process_ai_generated_images(
                requirement=requirement,
                project_topic=request.project_topic,
                project_scenario=request.project_scenario,
                slide_title=slide_content.get("title", ""),
                slide_content=slide_content.get("title", ""),
                image_config=image_config,
                page_number=request.slide_index + 1,
                total_pages=1,
                template_html=slide_content.get("html_content", ""),
            )
            images_collection.images.extend(ai_images)

        # 重新计算统计信息
        images_collection.__post_init__()

        if images_collection.total_count == 0:
            return {
                "success": False,
                "message": "未能生成任何图片，请检查配置和网络连接",
            }

        # 获取第一张生成的图像（用于替换）
        new_image = images_collection.images[0]
        new_image_url = new_image.absolute_url

        # 替换HTML中的图像
        updated_html = replace_image_in_html(
            slide_content.get("html_content", ""), image_info, new_image_url
        )

        logger.info(f"图片重新生成成功: {new_image.source.value}来源, URL: {new_image_url}")

        return {
            "success": True,
            "message": f"图像重新生成成功（来源：{new_image.source.value}）",
            "new_image_url": new_image_url,
            "new_image_id": new_image.image_id,
            "updated_html_content": updated_html,
            "generation_prompt": getattr(new_image, "generation_prompt", ""),
            "image_source": new_image.source.value,
            "ai_analysis": {
                "total_images_analyzed": 1,
                "reasoning": f"用户请求重新生成{image_context.get('image_purpose', '图像')}，选择{selected_source.value}来源",
                "enabled_sources": [source.value for source in enabled_sources],
                "selected_source": selected_source.value,
            },
            "image_info": {
                "width": new_image.width,
                "height": new_image.height,
                "format": getattr(new_image, "format", "unknown"),
                "alt_text": new_image.alt_text,
                "title": new_image.title,
                "source": new_image.source.value,
                "purpose": new_image.purpose.value,
            },
        }

    except Exception as e:
        logger.error(f"AI图像重新生成失败: {e}")
        import traceback

        traceback.print_exc()
        return {"success": False, "message": f"图像重新生成失败: {str(e)}"}


@router.post("/api/ai/auto-generate-slide-images")
async def ai_auto_generate_slide_images(
    request: AIAutoImageGenerateRequest, user: User = Depends(get_current_user_required)
):
    """AI一键配图接口 - 自动分析幻灯片内容并生成相关配图"""
    try:
        # 获取图像服务和AI提供者
        from ..services.image.image_service import get_image_service

        image_service = get_image_service()
        if not image_service:
            return {"success": False, "message": "图像服务不可用"}

        ai_provider = get_ai_provider()
        if not ai_provider:
            return {"success": False, "message": "AI提供者不可用"}

        # 获取图像处理器
        from ..services.ppt_image_processor import PPTImageProcessor

        image_processor = PPTImageProcessor(image_service, ai_provider)

        slide_content = request.slide_content
        slide_title = slide_content.get("title", f"第{request.slide_index + 1}页")
        slide_html = slide_content.get("html_content", "")

        logger.info(f"开始为第{request.slide_index + 1}页进行一键配图")

        # 第一步：AI分析幻灯片内容，确定是否需要配图以及配图需求
        analysis_prompt = f"""作为专业的Slide设计师，请分析以下幻灯片内容，判断是否需要配图以及配图需求。

项目主题：{request.project_topic}
项目场景：{request.project_scenario}
幻灯片标题：{slide_title}
幻灯片HTML内容：{slide_html[:1000]}...

请分析：
1. 这个幻灯片是否需要配图？
2. 如果需要，应该配几张图？
3. 每张图的用途和描述是什么？
4. 图片应该插入到什么位置？

请以JSON格式回复：
{{
    "needs_images": true/false,
    "image_count": 数量,
    "images": [
        {{
            "purpose": "图片用途（如：主要插图、装饰图、背景图等）",
            "description": "图片内容描述",
            "keywords": "搜索关键词",
            "position": "插入位置（如：标题下方、内容中间、页面右侧等）"
        }}
    ],
    "reasoning": "分析理由"
}}"""

        analysis_response = await ai_provider.text_completion(
            prompt=analysis_prompt, temperature=0.3
        )

        # 解析AI分析结果
        import json

        try:
            analysis_result = json.loads(analysis_response.content.strip())
        except json.JSONDecodeError:
            # 如果JSON解析失败，使用默认配置
            analysis_result = {
                "needs_images": True,
                "image_count": 1,
                "images": [
                    {
                        "purpose": "主要插图",
                        "description": f"与{slide_title}相关的配图",
                        "keywords": f"{request.project_topic} {slide_title}",
                        "position": "内容中间",
                    }
                ],
                "reasoning": "默认为幻灯片添加一张主要配图",
            }

        if not analysis_result.get("needs_images", False):
            return {
                "success": True,
                "message": "AI分析认为此幻灯片不需要配图",
                "updated_html_content": slide_html,
                "generated_images_count": 0,
                "ai_analysis": analysis_result,
            }

        # 第二步：根据分析结果生成图片需求
        from ..services.models.slide_image_info import (
            ImagePurpose,
            ImageRequirement,
            ImageSource,
            SlideImagesCollection,
        )

        images_collection = SlideImagesCollection(page_number=request.slide_index + 1, images=[])

        # 获取图像配置（使用与重新生成图片相同的配置键）
        from ..services.config_service import config_service

        image_config = config_service.get_config_by_category("image_service")

        # 检查是否启用图片生成服务
        enable_image_service = image_config.get("enable_image_service", False)
        if not enable_image_service:
            return {"success": False, "message": "图片生成服务未启用，请在配置中启用"}

    # 获取启用的图像来源（使用与重新生成图片相同的逻辑）

        enabled_sources = []
        if image_config.get("enable_local_images", True):
            enabled_sources.append(ImageSource.LOCAL)
        if image_config.get("enable_network_search", False):
            enabled_sources.append(ImageSource.NETWORK)
        if image_config.get("enable_ai_generation", False):
            enabled_sources.append(ImageSource.AI_GENERATED)

        if not enabled_sources:
            return {
                "success": False,
                "message": "没有启用的图像来源，请在设置中配置图像获取方式",
            }

        # 使用与重新生成图片完全相同的图片来源选择逻辑
        image_context = {
            "image_purpose": "illustration",  # 一键配图默认为说明性图片
            "slide_title": slide_title,
            "slide_content": slide_html,
        }

        selected_source = select_best_image_source(enabled_sources, image_config, image_context)

        # 为每个图片需求生成图片
        for i, image_info in enumerate(analysis_result.get("images", [])[:3]):  # 最多3张图
            # 创建图片需求
            requirement = ImageRequirement(
                purpose=ImagePurpose.ILLUSTRATION,
                description=image_info.get("description", "相关配图"),
                priority=1,
                source=selected_source,
                count=1,
            )

            # 根据选择的来源处理图片生成
            if (
                requirement.source == ImageSource.AI_GENERATED
                and ImageSource.AI_GENERATED in enabled_sources
            ):
                ai_images = await image_processor._process_ai_generated_images(
                    requirement=requirement,
                    project_topic=request.project_topic,
                    project_scenario=request.project_scenario,
                    slide_title=slide_title,
                    slide_content=slide_title,
                    image_config=image_config,
                    page_number=request.slide_index + 1,
                    total_pages=1,
                    template_html=slide_html,
                )
                images_collection.images.extend(ai_images)

            elif (
                requirement.source == ImageSource.NETWORK and ImageSource.NETWORK in enabled_sources
            ):
                network_images = await image_processor._process_network_images(
                    requirement=requirement,
                    project_topic=request.project_topic,
                    project_scenario=request.project_scenario,
                    slide_title=slide_title,
                    slide_content=slide_title,
                    image_config=image_config,
                )
                images_collection.images.extend(network_images)

            elif requirement.source == ImageSource.LOCAL and ImageSource.LOCAL in enabled_sources:
                local_images = await image_processor._process_local_images(
                    requirement=requirement,
                    project_topic=request.project_topic,
                    project_scenario=request.project_scenario,
                    slide_title=slide_title,
                    slide_content=slide_title,
                )
                images_collection.images.extend(local_images)

        if not images_collection.images:
            return {"success": False, "message": "未能生成任何配图，请检查图像服务配置"}

        # 第三步：将生成的图片插入到幻灯片中
        updated_html = await image_processor._insert_images_into_slide(
            slide_html, images_collection, slide_title
        )

        logger.info(f"一键配图完成: 生成{len(images_collection.images)}张图片")

        return {
            "success": True,
            "message": f"一键配图完成，已生成{len(images_collection.images)}张图片",
            "updated_html_content": updated_html,
            "generated_images_count": len(images_collection.images),
            "generated_images": [
                {
                    "image_id": img.image_id,
                    "url": img.absolute_url,
                    "description": img.content_description,
                    "source": img.source.value,
                }
                for img in images_collection.images
            ],
            "ai_analysis": analysis_result,
        }

    except Exception as e:
        logger.error(f"AI一键配图失败: {e}")
        return {"success": False, "message": f"一键配图失败: {str(e)}"}


@router.post("/api/ai/enhance-bullet-point")
async def ai_enhance_bullet_point(
    request: AIBulletPointEnhanceRequest,
    user: User = Depends(get_current_user_required),
):
    """AI增强要点接口"""
    try:
        # 获取AI提供者
        ai_provider = get_ai_provider()

        # 构建上下文信息
        context_info = ""
        if request.contextInfo:
            original_point = request.contextInfo.get("originalBulletPoint", "")
            other_points = request.contextInfo.get("otherBulletPoints", [])
            point_index = request.contextInfo.get("pointIndex", 0)

            context_info = f"""
当前要点上下文信息：
- 要点位置：第{point_index + 1}个要点
- 原始要点内容：{original_point}
- 同页面其他要点：{', '.join(other_points) if other_points else '无'}
"""

        # 构建大纲信息
        outline_info = ""
        if request.slideOutline:
            outline_info = f"""
当前幻灯片大纲信息：
- 幻灯片类型：{request.slideOutline.get('slide_type', '未知')}
- 描述：{request.slideOutline.get('description', '无')}
- 所有要点：{', '.join(request.slideOutline.get('content_points', [])) if request.slideOutline.get('content_points') else '无'}
"""

        # 构建AI增强提示词
        context = f"""
你是一位专业的Slide内容编辑专家。用户需要你增强和优化一个Slide要点的内容。

项目信息：
- 项目标题：{request.projectInfo.get('title', '未知')}
- 项目主题：{request.projectInfo.get('topic', '未知')}
- 应用场景：{request.projectInfo.get('scenario', '未知')}

幻灯片信息：
- 幻灯片标题：{request.slideTitle}
- 幻灯片位置：第{request.slideIndex}页

{outline_info}

{context_info}

用户请求：{request.userRequest}

请根据以上信息，对要点进行增强和优化。要求：

1. **保持核心意思不变**：不要改变要点的基本含义和方向
2. **增加具体细节**：添加更多具体的描述、数据、例子或说明
3. **提升表达质量**：使用更专业、更有吸引力的表达方式
4. **保持简洁性**：虽然要增强内容，但仍要保持要点的简洁特性，不要过于冗长
5. **与其他要点协调**：确保增强后的要点与同页面其他要点在风格和层次上保持一致
6. **符合场景需求**：根据应用场景调整语言风格和专业程度

请直接返回增强后的要点内容，不需要额外的解释或格式化。
"""

        # 调用AI生成增强内容
        response = await ai_provider.text_completion(
            prompt=context,
            max_tokens=ai_config.max_tokens // 2,  # 使用较小的token限制
            temperature=0.7,
        )

        enhanced_text = response.content.strip()

        # 简单的内容验证
        if not enhanced_text or len(enhanced_text) < 5:
            raise ValueError("AI生成的增强内容过短或为空")

        return {
            "success": True,
            "enhancedText": enhanced_text,
            "originalText": (
                request.contextInfo.get("originalBulletPoint", "") if request.contextInfo else ""
            ),
        }

    except Exception as e:
        logger.error(f"AI要点增强请求失败: {e}")
        return {
            "success": False,
            "error": str(e),
            "message": "抱歉，AI要点增强服务暂时不可用。请稍后重试。",
        }


@router.post("/api/ai/enhance-all-bullet-points")
async def ai_enhance_all_bullet_points(
    request: AIBulletPointEnhanceRequest,
    user: User = Depends(get_current_user_required),
):
    """AI增强所有要点接口"""
    try:
        # 获取AI提供者
        ai_provider = get_ai_provider()

        # 构建上下文信息
        context_info = ""
        all_points = []
        if request.contextInfo:
            all_points = request.contextInfo.get("allBulletPoints", [])
            total_points = request.contextInfo.get("totalPoints", 0)

            context_info = f"""
当前要点上下文信息：
- 要点总数：{total_points}个
- 所有要点内容：
"""
            for i, point in enumerate(all_points, 1):
                context_info += f"  {i}. {point}\n"

        # 构建大纲信息
        outline_info = ""
        if request.slideOutline:
            outline_info = f"""
当前幻灯片大纲信息：
- 幻灯片类型：{request.slideOutline.get('slide_type', '未知')}
- 描述：{request.slideOutline.get('description', '无')}
"""

        # 构建AI增强提示词
        context = f"""
请对以下Slide要点进行增强和优化。

项目背景：
- 项目：{request.projectInfo.get('title', '未知')}
- 主题：{request.projectInfo.get('topic', '未知')}
- 场景：{request.projectInfo.get('scenario', '未知')}
- 幻灯片：{request.slideTitle}（第{request.slideIndex}页）

{outline_info}

{context_info}

增强要求：
1. 保持每个要点的核心意思不变
2. 添加具体细节、数据或例子
3. 使用更专业、准确的表达
4. 保持简洁，避免冗长
5. 确保要点间逻辑连贯、风格统一
6. 符合{request.projectInfo.get('scenario', '商务')}场景的专业要求

重要：请直接返回增强后的要点列表，每行一个要点，不要包含任何解释、开场白或格式说明。不要添加编号、符号或其他标记。

示例格式：
第一个增强后的要点内容
第二个增强后的要点内容
第三个增强后的要点内容
"""

        # 调用AI生成增强内容
        response = await ai_provider.text_completion(
            prompt=context,
            max_tokens=ai_config.max_tokens,  # 使用完整的token限制，因为要处理多个要点
            temperature=0.7,
        )

        enhanced_content = response.content.strip()

        # 解析增强后的要点 - 改进的过滤逻辑
        enhanced_points = []
        if enhanced_content:
            # 按行分割，过滤空行
            lines = [line.strip() for line in enhanced_content.split("\n") if line.strip()]

            # 过滤掉常见的无关内容
            filtered_lines = []
            skip_patterns = [
                "好的",
                "作为",
                "我将",
                "我会",
                "以下是",
                "根据",
                "请注意",
                "需要说明",
                "增强后的要点",
                "优化后的",
                "改进后的",
                "以上",
                "总结",
                "希望",
                "如有",
                "如果",
                "建议",
                "推荐",
                "注意",
                "提醒",
                "说明",
                "要点1",
                "要点2",
                "要点3",
                "要点4",
                "要点5",
                "第一",
                "第二",
                "第三",
                "第四",
                "第五",
                "第六",
                "第七",
                "第八",
                "第九",
                "第十",
                "1.",
                "2.",
                "3.",
                "4.",
                "5.",
                "6.",
                "7.",
                "8.",
                "9.",
                "10.",
                "•",
                "·",
                "-",
                "*",
                "→",
                "▪",
                "▫",
            ]

            for line in lines:
                # 跳过过短的行（可能是格式标记）
                if len(line) < 5:
                    continue

                # 跳过包含常见开场白模式的行
                should_skip = False
                for pattern in skip_patterns:
                    if line.startswith(pattern) or (
                        pattern in ["好的", "作为", "我将", "我会"] and pattern in line[:10]
                    ):
                        should_skip = True
                        break

                # 跳过纯数字或符号开头的行（可能是编号）
                if line[0].isdigit() or line[0] in ["•", "·", "-", "*", "→", "▪", "▫"]:
                    # 但保留去掉编号后的内容
                    cleaned_line = line
                    # 移除开头的编号和符号
                    import re

                    cleaned_line = re.sub(r"^[\d\s\.\-\*\•\·\→\▪\▫]+", "", cleaned_line).strip()
                    if len(cleaned_line) >= 5:
                        filtered_lines.append(cleaned_line)
                    continue

                if not should_skip:
                    filtered_lines.append(line)

            enhanced_points = filtered_lines

        # 简单的内容验证
        if not enhanced_points or len(enhanced_points) == 0:
            raise ValueError("AI生成的增强内容为空或被过滤")

        return {
            "success": True,
            "enhancedPoints": enhanced_points,
            "originalPoints": all_points,
            "totalEnhanced": len(enhanced_points),
        }

    except Exception as e:
        logger.error(f"AI增强所有要点请求失败: {e}")
        return {
            "success": False,
            "error": str(e),
            "message": "抱歉，AI要点增强服务暂时不可用。请稍后重试。",
        }


@router.post("/api/ai/simple_completion")
async def api_simple_completion(
    request: SimpleCompletionRequest,
    user: User = Depends(get_current_user_required),
):
    """轻量级通用文本补全/润色接口。

    前端仅需提供 prompt 字段。返回：{ success, content }
    - 默认使用较低温度（0.55）保证润色稳定性
    - max_tokens 默认使用全局配置 1/2，用户可覆盖
    """
    try:
        if not request.prompt or len(request.prompt.strip()) < 4:
            return {"success": False, "error": "提示词过短"}

        ai_provider = get_ai_provider()
        max_tokens = request.max_tokens or max(256, ai_config.max_tokens // 2)
        temperature = 0.55 if request.temperature is None else request.temperature

        # 调用底层 completion
        resp = await ai_provider.text_completion(
            prompt=request.prompt.strip(),
            max_tokens=max_tokens,
            temperature=temperature,
        )

        content = (resp.content or "").strip()

        # 简单清洗：去除典型开场语句
        if content.startswith("当然") or content.startswith("好的"):
            # 只移除首句的这些寒暄，避免删过多
            parts = re.split(r"[。.!？!?]", content, 1)
            if len(parts) == 2:
                content = parts[1].strip()

        if not content:
            return {"success": False, "error": "AI返回为空"}

        return {"success": True, "content": content}
    except Exception as e:
        logger.error(f"simple_completion 失败: {e}")
        return {"success": False, "error": str(e)}


@router.get("/api/projects/{project_id}/selected-global-template")
async def get_selected_global_template(
    project_id: str, user: User = Depends(get_current_user_required)
):
    """获取项目选择的全局母版模板"""
    try:
        # Access control
        project = await ppt_service.project_manager.get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        if not _user_can_access_project(user, project):
            raise HTTPException(status_code=403, detail="Access denied")

        # 检查项目是否真正选择了模板
        selected_template = await ppt_service.get_selected_global_template(project_id)
        if selected_template:
            logger.info(
                f"Project {project_id} has selected template: {selected_template.get('template_name', 'Unknown')}"
            )
            return {
                "status": "success",
                "template": selected_template,
                "is_user_selected": True,
            }
        else:
            # 如果没有选择的模板，尝试获取默认模板
            default_template = await ppt_service.global_template_service.get_default_template()
            if default_template:
                logger.info(
                    f"Project {project_id} using default template: {default_template.get('template_name', 'Unknown')}"
                )
                return {
                    "status": "success",
                    "template": default_template,
                    "is_user_selected": False,
                }
            else:
                logger.warning(f"No template available for project {project_id}")
                return {
                    "status": "success",
                    "template": None,
                    "is_user_selected": False,
                }
    except Exception as e:
        logger.error(f"Error getting selected global template for project {project_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/api/projects/{project_id}/slides/{slide_index}/save")
async def save_single_slide_content(
    project_id: str,
    slide_index: int,
    request: Request,
    user: User = Depends(get_current_user_required),
):
    """保存单个幻灯片内容到数据库"""
    try:
        logger.info(f"🔄 开始保存项目 {project_id} 的第 {slide_index + 1} 页 (索引: {slide_index})")

        data = await request.json()
        html_content = data.get("html_content", "")

        logger.info(f"📄 接收到HTML内容，长度: {len(html_content)} 字符")

        if not html_content:
            logger.error("❌ HTML内容为空")
            raise HTTPException(status_code=400, detail="HTML content is required")

        project = await ppt_service.project_manager.get_project(project_id)
        if not project:
            logger.error(f"❌ 项目 {project_id} 不存在")
            raise HTTPException(status_code=404, detail="Project not found")
        if not _user_can_access_project(user, project):
            raise HTTPException(status_code=403, detail="Access denied")

        # 详细验证幻灯片索引
        total_slides = len(project.slides_data) if project.slides_data else 0
        logger.debug(f"📊 项目幻灯片信息: 总页数={total_slides}, 请求索引={slide_index}")

        if slide_index < 0:
            logger.error(f"❌ 幻灯片索引不能为负数: {slide_index}")
            raise HTTPException(
                status_code=400, detail=f"Slide index cannot be negative: {slide_index}"
            )

        if slide_index >= total_slides:
            logger.error(f"❌ 幻灯片索引超出范围: {slide_index}，项目共有 {total_slides} 页")
            raise HTTPException(
                status_code=400,
                detail=f"Slide index {slide_index} out of range (total: {total_slides})",
            )

        logger.debug(f"📝 更新第 {slide_index + 1} 页的内容")

        # 更新幻灯片数据
        project.slides_data[slide_index]["html_content"] = html_content
        project.slides_data[slide_index]["is_user_edited"] = True
        project.updated_at = time.time()

        # 重新生成组合HTML
        if project.slides_data:
            outline_title = project.title
            if isinstance(project.outline, dict):
                outline_title = project.outline.get("title", project.title)
            elif hasattr(project.outline, "title"):
                outline_title = project.outline.title

            project.slides_html = ppt_service._combine_slides_to_full_html(
                project.slides_data, outline_title
            )

        # 保存到数据库
        try:
            logger.debug(f"💾 开始保存到数据库... (第{slide_index + 1}页)")

            from ..services.db_project_manager import DatabaseProjectManager

            db_manager = DatabaseProjectManager()

            # 保存单个幻灯片
            slide_data = project.slides_data[slide_index]
            slide_title = slide_data.get("title", "无标题")
            is_user_edited = slide_data.get("is_user_edited", False)

            logger.debug(
                f"📊 幻灯片数据: 标题='{slide_title}', 用户编辑={is_user_edited}, 索引={slide_index}"
            )
            logger.debug(f"🔍 保存前验证: 项目ID={project_id}, 幻灯片索引={slide_index}")

            save_success = await db_manager.save_single_slide(project_id, slide_index, slide_data)

            if save_success:
                logger.debug(f"✅ 第 {slide_index + 1} 页已成功保存到数据库")

                return {
                    "success": True,
                    "message": f"Slide {slide_index + 1} saved successfully to database",
                    "slide_data": slide_data,
                    "database_saved": True,
                }
            else:
                logger.error(f"❌ 保存第 {slide_index + 1} 页到数据库失败")
                return {
                    "success": False,
                    "error": "Failed to save slide to database",
                    "database_saved": False,
                }

        except Exception as save_error:
            logger.error(f"❌ 保存第 {slide_index + 1} 页时发生异常: {save_error}")
            import traceback

            traceback.print_exc()
            return {
                "success": False,
                "error": f"Database error: {str(save_error)}",
                "database_saved": False,
            }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ 保存单个幻灯片时发生错误: {e}")
        import traceback

        traceback.print_exc()
        return {"success": False, "error": str(e), "database_saved": False}


@router.get("/api/projects/{project_id}/slides/stream")
async def stream_slides_generation(
    project_id: str, user: User = Depends(get_current_user_required)
):
    """Stream slides generation process"""
    try:
        # Access control
        project = await ppt_service.project_manager.get_project(project_id)
        if not project:
            return {"error": "Project not found"}
        if not _user_can_access_project(user, project):
            return {"error": "Access denied"}

        async def generate_slides_stream():
            async for chunk in ppt_service.generate_slides_streaming(project_id):
                yield chunk

        return StreamingResponse(
            generate_slides_stream(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Headers": "Cache-Control",
            },
        )

    except Exception as e:
        return {"error": str(e)}


@router.post("/api/projects/{project_id}/slides/cleanup")
async def cleanup_excess_slides(
    project_id: str, request: Request, user: User = Depends(get_current_user_required)
):
    """清理项目中多余的幻灯片"""
    try:
        logger.info(f"🧹 开始清理项目 {project_id} 的多余幻灯片")

        data = await request.json()
        current_slide_count = data.get("current_slide_count", 0)

        if current_slide_count <= 0:
            logger.error("❌ 无效的幻灯片数量")
            raise HTTPException(status_code=400, detail="Invalid slide count")

        project = await ppt_service.project_manager.get_project(project_id)
        if not project:
            logger.error(f"❌ 项目 {project_id} 不存在")
            raise HTTPException(status_code=404, detail="Project not found")
        if not _user_can_access_project(user, project):
            raise HTTPException(status_code=403, detail="Access denied")

        # 清理数据库中多余的幻灯片
        from ..services.db_project_manager import DatabaseProjectManager

        db_manager = DatabaseProjectManager()
        deleted_count = await db_manager.cleanup_excess_slides(project_id, current_slide_count)

        logger.info(f"✅ 项目 {project_id} 清理完成，删除了 {deleted_count} 张多余的幻灯片")

        return {
            "success": True,
            "message": f"Successfully cleaned up {deleted_count} excess slides",
            "deleted_count": deleted_count,
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ 清理幻灯片失败: {e}")
        return {"success": False, "error": str(e)}


@router.post("/api/projects/{project_id}/slides/batch-save")
async def batch_save_slides(
    project_id: str, request: Request, user: User = Depends(get_current_user_required)
):
    """批量保存所有幻灯片 - 高效版本"""
    try:
        logger.debug(f"🔄 开始批量保存项目 {project_id} 的所有幻灯片")

        data = await request.json()
        slides_data = data.get("slides_data", [])

        if not slides_data:
            logger.error("❌ 幻灯片数据为空")
            raise HTTPException(status_code=400, detail="Slides data is required")

        project = await ppt_service.project_manager.get_project(project_id)
        if not project:
            logger.error(f"❌ 项目 {project_id} 不存在")
            raise HTTPException(status_code=404, detail="Project not found")
        if not _user_can_access_project(user, project):
            raise HTTPException(status_code=403, detail="Access denied")

        # 更新项目内存中的数据
        project.slides_data = slides_data
        project.updated_at = time.time()

        # 重新生成完整HTML
        outline_title = project.title
        if hasattr(project, "outline") and project.outline:
            outline_title = project.outline.get("title", project.title)

        project.slides_html = ppt_service._combine_slides_to_full_html(
            project.slides_data, outline_title
        )

        # 使用批量保存到数据库
        from ..services.db_project_manager import DatabaseProjectManager

        db_manager = DatabaseProjectManager()

        # 批量保存幻灯片
        batch_success = await db_manager.batch_save_slides(project_id, slides_data)

        # 更新项目信息
        if batch_success:
            await db_manager.update_project_data(
                project_id,
                {
                    "slides_html": project.slides_html,
                    "slides_data": project.slides_data,
                    "updated_at": project.updated_at,
                },
            )

        logger.debug(f"✅ 项目 {project_id} 批量保存完成，共 {len(slides_data)} 张幻灯片")

        return {
            "success": batch_success,
            "message": (
                f"Successfully batch saved {len(slides_data)} slides"
                if batch_success
                else "Batch save failed"
            ),
            "slides_count": len(slides_data),
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ 批量保存幻灯片失败: {e}")
        return {"success": False, "error": str(e)}


@router.get("/api/test-pdf-export")
async def test_pdf_export_endpoint():
    """Test endpoint to verify PDF export is working"""
    return {"status": "PDF export endpoint is working", "timestamp": time.time()}


@router.get("/api/projects/{project_id}/export/pdf")
async def export_project_pdf(
    project_id: str, individual: bool = False, user: User = Depends(get_current_user_optional)
):
    """Export project as PDF using Pyppeteer"""
    print(f"🔥 PDF EXPORT REQUEST RECEIVED: project_id={project_id}")
    logging.info(f"🔥 PDF EXPORT REQUEST RECEIVED: project_id={project_id}")
    try:
        project = await ppt_service.project_manager.get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        if user and not _user_can_access_project(user, project):
            raise HTTPException(status_code=403, detail="Access denied")

        # Check if we have slides data
        if not project.slides_data or len(project.slides_data) == 0:
            raise HTTPException(status_code=400, detail="PPT not generated yet")

        # Check if Pyppeteer is available
        pdf_converter = get_pdf_converter()
        if not pdf_converter.is_available():
            raise HTTPException(
                status_code=503,
                detail="PDF generation service unavailable. Please ensure Pyppeteer is installed: pip install pyppeteer",
            )

        # Create temp file in thread pool to avoid blocking
        temp_pdf_path = await run_blocking_io(
            lambda: tempfile.NamedTemporaryFile(suffix=".pdf", delete=False).name
        )

        logging.info(f"Generating PDF for project: {project.topic}")
        logging.info(f"Project has {len(project.slides_data)} slides")

        # Prefer simple, single-pass path: one combined HTML rendered once
        success = False
        try:
            combined_html = await _generate_combined_html_for_pdf(project)
            # Use converter directly in simple mode
            success = await pdf_converter.convert_single_html_to_pdf(
                combined_html, temp_pdf_path, simple=True
            )
            if success:
                logging.info("PDF generated via single-pass simple mode")
        except Exception as e:
            logging.info(f"Simple single-pass PDF attempt failed, will fallback: {e}")

        if not success:
            # Fallback: legacy per-slide HTML -> batch convert -> merge
            success = await _generate_pdf_with_pyppeteer(project, temp_pdf_path, individual)

        if not success:
            # Clean up temp file and raise error
            logging.error(f"❌ PDF generation failed for project {project.project_id}")
            print(f"❌ PDF generation failed for project {project.project_id}")
            await run_blocking_io(
                lambda: (os.unlink(temp_pdf_path) if os.path.exists(temp_pdf_path) else None)
            )
            raise HTTPException(status_code=500, detail="PDF generation failed")

        # Return PDF file
        logging.info("PDF generated successfully")

        # 创建安全的文件名（ASCII 回退 + UTF-8 显示名）
        safe_topic = re.sub(r"[^\w\s-]", "", project.topic).strip()[:50] or "presentation"
        ascii_topic = re.sub(r"[^A-Za-z0-9_-]", "", safe_topic) or "presentation"
        ascii_filename = f"{ascii_topic}_PPT.pdf"
        utf8_filename = f"{safe_topic}_PPT.pdf"
        encoded_filename = urllib.parse.quote(utf8_filename, safe="")

        # 使用BackgroundTask来清理临时文件
        from starlette.background import BackgroundTask

        def cleanup_temp_file():
            try:
                os.unlink(temp_pdf_path)
            except Exception:
                pass

        return FileResponse(
            temp_pdf_path,
            media_type="application/pdf",
            headers={
                # ASCII 回退 + UTF-8 filename*
                "Content-Disposition": f"attachment; filename=\"{ascii_filename}\"; filename*=UTF-8''{encoded_filename}",
                "X-PDF-Generator": "Pyppeteer",
            },
            background=BackgroundTask(cleanup_temp_file),
        )

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"❌ PDF export exception: {e}", exc_info=True)
        print(f"❌ PDF export exception: {e}")
        import traceback

        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"PDF export failed: {str(e)}")


@router.get("/api/projects/{project_id}/export/pdf/individual")
async def export_project_pdf_individual(
    project_id: str, user: User = Depends(get_current_user_required)
):
    """Export project as individual PDF files for each slide"""
    return await export_project_pdf(project_id, individual=True, user=user)


@router.get("/api/apryse/license/check")
async def check_apryse_license(user: User = Depends(get_current_user_required)):
    """Check if Apryse license is valid and SDK is available"""
    try:
        converter = get_pdf_to_pptx_converter()

        # Check if SDK is available
        if not converter._check_sdk_availability():
            return {
                "valid": False,
                "error": "Apryse SDK not available. Please ensure it's installed: pip install apryse-sdk",
                "details": "SDK files or Python package not found",
            }

        # Check if license can initialize SDK
        if not converter._initialize_sdk():
            return {
                "valid": False,
                "error": "License key is invalid or SDK initialization failed",
                "details": "Please check your license key and ensure it's valid",
            }

        return {
            "valid": True,
            "message": "Apryse license is valid and SDK is ready",
            "details": "PPTX conversion is available",
        }

    except Exception as e:
        return {
            "valid": False,
            "error": f"License check failed: {str(e)}",
            "details": "Unexpected error during license validation",
        }


@router.get("/api/projects/{project_id}/export/pptx")
async def export_project_pptx(project_id: str, user: User = Depends(get_current_user_optional)):
    """Export project as PPTX by first generating PDF then converting to PowerPoint"""
    try:
        project = await ppt_service.project_manager.get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        if user and not _user_can_access_project(user, project):
            raise HTTPException(status_code=403, detail="Access denied")

        # Check if we have slides data
        if not project.slides_data or len(project.slides_data) == 0:
            raise HTTPException(status_code=400, detail="PPT not generated yet")

        # Get PDF to PPTX converter
        converter = get_pdf_to_pptx_converter()
        if not converter.is_available():
            raise HTTPException(
                status_code=503,
                detail="PPTX conversion service unavailable. Please ensure Apryse SDK is installed and licensed.",
            )

        # Check if Pyppeteer is available for PDF generation
        pdf_converter = get_pdf_converter()
        if not pdf_converter.is_available():
            raise HTTPException(
                status_code=503,
                detail="PDF generation service unavailable. Please ensure Pyppeteer is installed: pip install pyppeteer",
            )

        # Step 1: Generate PDF using existing PDF export functionality
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as temp_pdf_file:
            temp_pdf_path = temp_pdf_file.name

        logging.info("Step 1: Generating PDF for PPTX conversion")
        pdf_success = await _generate_pdf_with_pyppeteer(project, temp_pdf_path, individual=False)

        if not pdf_success:
            # Clean up temp file and raise error
            try:
                os.unlink(temp_pdf_path)
            except Exception:
                pass
            raise HTTPException(status_code=500, detail="PDF generation failed")

        # Step 2: Convert PDF to PPTX
        logging.info("Step 2: Converting PDF to PPTX")

        # Create temporary PPTX file
        with tempfile.NamedTemporaryFile(suffix=".pptx", delete=False) as temp_pptx_file:
            temp_pptx_path = temp_pptx_file.name

        try:
            # Perform PDF to PPTX conversion in thread pool to avoid blocking
            conversion_success, result_path = await run_blocking_io(
                converter.convert_pdf_to_pptx, temp_pdf_path, temp_pptx_path
            )

            if not conversion_success:
                raise HTTPException(
                    status_code=500,
                    detail=f"PDF to PPTX conversion failed: {result_path}",
                )

            # Verify PPTX file was created successfully
            if not os.path.exists(temp_pptx_path) or os.path.getsize(temp_pptx_path) == 0:
                raise HTTPException(status_code=500, detail="PPTX file was not created or is empty")

            # Return PPTX file
            logging.info("PPTX generated successfully")

            # 创建安全的文件名（ASCII 回退 + UTF-8 显示名）
            safe_topic = re.sub(r"[^\w\s-]", "", project.topic).strip()[:50] or "presentation"
            ascii_topic = re.sub(r"[^A-Za-z0-9_-]", "", safe_topic) or "presentation"
            ascii_filename = f"{ascii_topic}_PPT.pptx"
            utf8_filename = f"{safe_topic}_PPT.pptx"
            encoded_filename = urllib.parse.quote(utf8_filename, safe="")

            # 使用BackgroundTask来清理临时文件
            from starlette.background import BackgroundTask

            def cleanup_temp_files():
                try:
                    os.unlink(temp_pdf_path)
                except Exception:
                    pass
                try:
                    os.unlink(temp_pptx_path)
                except Exception:
                    pass

            return FileResponse(
                temp_pptx_path,
                media_type="application/vnd.openxmlformats-officedocument.presentationml.presentation",
                headers={
                    # ASCII 回退 + UTF-8 filename*
                    "Content-Disposition": f"attachment; filename=\"{ascii_filename}\"; filename*=UTF-8''{encoded_filename}",
                    "X-Conversion-Method": "PDF-to-PPTX",
                },
                background=BackgroundTask(cleanup_temp_files),
            )

        except Exception as e:
            # Clean up temp files on error
            try:
                os.unlink(temp_pdf_path)
            except Exception:
                pass
            try:
                os.unlink(temp_pptx_path)
            except Exception:
                pass
            raise HTTPException(status_code=500, detail=f"PPTX conversion error: {str(e)}")

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/api/projects/{project_id}/export/html")
async def export_project_html(project_id: str, user: User = Depends(get_current_user_required)):
    """Export project as HTML ZIP package with slideshow index"""
    try:
        project = await ppt_service.project_manager.get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        if not _user_can_access_project(user, project):
            raise HTTPException(status_code=403, detail="Access denied")

        # Check if we have slides data
        if not project.slides_data or len(project.slides_data) == 0:
            raise HTTPException(status_code=400, detail="Slide not generated yet")

        # Create temporary directory and generate files in thread pool
        zip_content = await run_blocking_io(_generate_html_export_sync, project)

        # URL encode the filename to handle Chinese characters
        zip_filename = f"{project.topic}_Slide.zip"
        safe_filename = urllib.parse.quote(zip_filename, safe="")

        from fastapi.responses import Response

        return Response(
            content=zip_content,
            media_type="application/zip",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_filename}"},
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


def _generate_html_export_sync(project) -> bytes:
    """同步生成HTML导出文件（在线程池中运行）"""
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)

        # Generate individual HTML files for each slide
        slide_files = []
        for i, slide in enumerate(project.slides_data):
            slide_filename = f"slide_{i+1}.html"
            slide_files.append(slide_filename)

            # Create complete HTML document for each slide
            slide_html = _generate_individual_slide_html_sync(
                slide, i + 1, len(project.slides_data), project.topic
            )

            slide_path = temp_path / slide_filename
            with open(slide_path, "w", encoding="utf-8") as f:
                f.write(slide_html)

        # Generate index.html slideshow page
        index_html = _generate_slideshow_index_sync(project, slide_files)
        index_path = temp_path / "index.html"
        with open(index_path, "w", encoding="utf-8") as f:
            f.write(index_html)

        # Create ZIP file
        zip_filename = f"{project.topic}_Slide.zip"
        zip_path = temp_path / zip_filename

        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
            # Add index.html
            zipf.write(index_path, "index.html")

            # Add all slide files
            for slide_file in slide_files:
                slide_path = temp_path / slide_file
                zipf.write(slide_path, slide_file)

        # Read ZIP file content
        with open(zip_path, "rb") as f:
            return f.read()


def _generate_individual_slide_html_sync(
    slide, slide_number: int, total_slides: int, topic: str
) -> str:
    """同步生成单个幻灯片HTML（在线程池中运行）"""
    slide_html = slide.get("html_content", "")
    slide_title = slide.get("title", f"第{slide_number}页")

    # Check if it's already a complete HTML document

    if slide_html.strip().lower().startswith("<!doctype") or slide_html.strip().lower().startswith(
        "<html"
    ):
        # It's a complete HTML document, enhance it with navigation
        return _enhance_complete_html_with_navigation(
            slide_html, slide_number, total_slides, topic, slide_title
        )
    else:
        # It's just content, wrap it in a complete structure
        slide_content = slide_html

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{topic} - {slide_title}</title>
    <style>
        body {{
            margin: 0;
            padding: 0;
            font-family: 'Microsoft YaHei', 'PingFang SC', sans-serif;
            background: #f5f5f5;
            display: flex;
            align-items: center;
            justify-content: center;
            min-height: 100vh;
        }}
        .slide-container {{
            width: 90vw;
            height: 90vh;
            background: white;
            border-radius: 10px;
            box-shadow: 0 4px 20px rgba(0,0,0,0.1);
            overflow: hidden;
            position: relative;
        }}
        .slide-content {{
            width: 100%;
            height: 100%;
            padding: 20px;
            box-sizing: border-box;
        }}
        .slide-number {{
            position: absolute;
            bottom: 20px;
            right: 20px;
            background: rgba(0,0,0,0.7);
            color: white;
            padding: 5px 10px;
            border-radius: 5px;
            font-size: 14px;
        }}
    </style>
</head>
<body>
    <div class="slide-container">
        <div class="slide-content">
            {slide_content}
        </div>
        <div class="slide-number">{slide_number} / {total_slides}</div>
    </div>
</body>
</html>"""


def _generate_slideshow_index_sync(project, slide_files: list) -> str:
    """同步生成幻灯片索引页面（在线程池中运行）"""
    slides_list = ""
    for i, slide_file in enumerate(slide_files):
        slide = project.slides_data[i]
        slide_title = slide.get("title", f"第{i+1}页")
        slides_list += f"""
        <div class="slide-item" onclick="openSlide('{slide_file}')">
            <div class="slide-preview">
                <div class="slide-number">{i+1}</div>
                <div class="slide-title">{slide_title}</div>
            </div>
        </div>"""

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{project.topic} - Slide放映</title>
    <style>
        body {{
            margin: 0;
            padding: 0;
            font-family: 'Microsoft YaHei', 'PingFang SC', sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
        }}
        .header {{
            text-align: center;
            padding: 40px 20px;
            color: white;
        }}
        .header h1 {{
            margin: 0;
            font-size: 2.5em;
            font-weight: 300;
        }}
        .slides-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fill, minmax(300px, 1fr));
            gap: 20px;
            padding: 20px;
        }}
        .slide-item {{
            background: white;
            border-radius: 10px;
            padding: 20px;
            cursor: pointer;
            transition: all 0.3s ease;
            box-shadow: 0 4px 15px rgba(0,0,0,0.1);
        }}
        .slide-item:hover {{
            transform: translateY(-5px);
            box-shadow: 0 8px 25px rgba(0,0,0,0.2);
        }}
        .slide-number {{
            background: #007bff;
            color: white;
            width: 40px;
            height: 40px;
            border-radius: 50%;
            display: flex;
            align-items: center;
            justify-content: center;
            margin: 0 auto 15px auto;
            font-weight: bold;
        }}
        .slide-title {{
            font-size: 1.1em;
            color: #333;
            margin: 0;
            text-align: center;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{project.topic}</h1>
        <p>Slide演示文稿 - 共{len(slide_files)}页</p>
    </div>
    <div class="slides-grid">
        {slides_list}
    </div>
    <script>
        function openSlide(slideFile) {{
            window.open(slideFile, '_blank');
        }}
    </script>
</body>
</html>"""


async def _generate_combined_html_for_export(project, export_type: str) -> str:
    """Generate combined HTML for export (PDF or HTML)"""
    try:
        if not project.slides_data:
            raise ValueError("No slides data available")

        # Create a combined HTML document with all slides
        html_parts = []

        # HTML document header
        html_parts.append(
            f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{project.topic} - Slide导出</title>
    <style>
        body {{
            margin: 0;
            padding: 0;
            font-family: 'Microsoft YaHei', 'PingFang SC', sans-serif;
            background: #f5f5f5;
        }}
        .slide-container {{
            width: 100vw;
            height: 100vh;
            display: flex;
            align-items: center;
            justify-content: center;
            page-break-after: always;
            background: white;
            position: relative;
        }}
        .slide-container:last-child {{
            page-break-after: avoid;
        }}
        .slide-frame {{
            width: 90vw;
            height: 90vh;
            border: none;
            border-radius: 10px;
            box-shadow: 0 4px 20px rgba(0,0,0,0.1);
        }}
        .slide-number {{
            position: absolute;
            bottom: 20px;
            right: 20px;
            background: rgba(0,0,0,0.7);
            color: white;
            padding: 5px 10px;
            border-radius: 5px;
            font-size: 14px;
        }}
        @media print {{
            .slide-container {{
                page-break-after: always;
                width: 100%;
                height: 100vh;
            }}
        }}
    </style>
</head>
<body>"""
        )

        # Add each slide preserving original styles
        for i, slide in enumerate(project.slides_data):
            slide_html = slide.get("html_content", "")
            if slide_html:
                # Preserve complete HTML structure
                if slide_html.strip().lower().startswith(
                    "<!doctype"
                ) or slide_html.strip().lower().startswith("<html"):
                    # Extract styles from head and content from body
                    import re

                    # Extract CSS styles from head
                    style_matches = re.findall(
                        r"<style[^>]*>(.*?)</style>",
                        slide_html,
                        re.DOTALL | re.IGNORECASE,
                    )
                    slide_styles = "\n".join(style_matches)

                    # Extract body content
                    body_match = re.search(
                        r"<body[^>]*>(.*?)</body>",
                        slide_html,
                        re.DOTALL | re.IGNORECASE,
                    )
                    if body_match:
                        slide_content = body_match.group(1)
                    else:
                        slide_content = slide_html
                else:
                    slide_styles = ""
                    slide_content = slide_html

                html_parts.append(
                    f"""
    <div class="slide-container">
        <style>
            {slide_styles}
        </style>
        <div class="slide-frame">
            {slide_content}
        </div>
        <div class="slide-number">{i + 1} / {len(project.slides_data)}</div>
    </div>"""
                )

        # Close HTML document
        html_parts.append(
            """
</body>
</html>"""
        )

        return "".join(html_parts)

    except Exception as e:
        # Fallback: return a simple error page
        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <title>导出错误</title>
</head>
<body>
    <h1>导出失败</h1>
    <p>错误信息: {str(e)}</p>
    <p>请确保Slide已经生成完成后再尝试导出。</p>
</body>
</html>"""


# Legacy Node.js Puppeteer check function - no longer needed with Pyppeteer
# def _check_puppeteer_available() -> bool:
#     """Check if Node.js and Puppeteer are available"""
#     # This function is deprecated - we now use Pyppeteer (Python) instead
#     return False


async def _generate_individual_slide_html(
    slide, slide_number: int, total_slides: int, topic: str
) -> str:
    """Generate complete HTML document for individual slide preserving original styles"""
    slide_html = slide.get("html_content", "")
    slide_title = slide.get("title", f"第{slide_number}页")

    # Check if it's already a complete HTML document

    if slide_html.strip().lower().startswith("<!doctype") or slide_html.strip().lower().startswith(
        "<html"
    ):
        # It's a complete HTML document, enhance it with navigation
        return _enhance_complete_html_with_navigation(
            slide_html, slide_number, total_slides, topic, slide_title
        )
    else:
        # It's just content, wrap it in a complete structure
        slide_content = slide_html

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{topic} - {slide_title}</title>
    <style>
        body {{
            margin: 0;
            padding: 0;
            font-family: 'Microsoft YaHei', 'PingFang SC', sans-serif;
            background: #f5f5f5;
            display: flex;
            align-items: center;
            justify-content: center;
            min-height: 100vh;
        }}
        .slide-container {{
            width: 90vw;
            height: 90vh;
            background: white;
            border-radius: 10px;
            box-shadow: 0 4px 20px rgba(0,0,0,0.1);
            overflow: hidden;
            position: relative;
        }}
        .slide-content {{
            width: 100%;
            height: 100%;
            padding: 20px;
            box-sizing: border-box;
        }}
        .slide-number {{
            position: absolute;
            bottom: 20px;
            right: 20px;
            background: rgba(0,0,0,0.7);
            color: white;
            padding: 5px 10px;
            border-radius: 5px;
            font-size: 14px;
        }}
        .navigation {{
            position: fixed;
            bottom: 20px;
            left: 50%;
            transform: translateX(-50%);
            display: flex;
            gap: 10px;
            z-index: 1000;
        }}
        .nav-btn {{
            background: #007bff;
            color: white;
            border: none;
            padding: 10px 15px;
            border-radius: 5px;
            cursor: pointer;
            text-decoration: none;
            display: inline-block;
        }}
        .nav-btn:hover {{
            background: #0056b3;
        }}
        .nav-btn:disabled {{
            background: #ccc;
            cursor: not-allowed;
        }}
        .fullscreen-btn {{
            position: fixed;
            top: 20px;
            right: 20px;
            background: #28a745;
            color: white;
            border: none;
            padding: 10px;
            border-radius: 5px;
            cursor: pointer;
            z-index: 1000;
        }}
        .fullscreen-btn:hover {{
            background: #1e7e34;
        }}
    </style>
</head>
<body>
    <div class="slide-container">
        <div class="slide-content">
            {slide_content}
        </div>
        <div class="slide-number">{slide_number} / {total_slides}</div>
    </div>

    <div class="navigation">
        <a href="index.html" class="nav-btn">🏠 返回目录</a>
        {"" if slide_number <= 1 else f'<a href="slide_{slide_number-1}.html" class="nav-btn">‹ 上一页</a>'}
        {"" if slide_number >= total_slides else f'<a href="slide_{slide_number+1}.html" class="nav-btn">下一页 ›</a>'}
    </div>

    <button class="fullscreen-btn" onclick="toggleFullscreen()" title="全屏显示">
        📺
    </button>

    <script>
        function toggleFullscreen() {{
            if (!document.fullscreenElement) {{
                document.documentElement.requestFullscreen();
            }} else {{
                if (document.exitFullscreen) {{
                    document.exitFullscreen();
                }}
            }}
        }}

        // Keyboard navigation
        document.addEventListener('keydown', function(e) {{
            if (e.key === 'ArrowLeft' && {slide_number} > 1) {{
                window.location.href = 'slide_{slide_number-1}.html';
            }} else if (e.key === 'ArrowRight' && {slide_number} < {total_slides}) {{
                window.location.href = 'slide_{slide_number+1}.html';
            }} else if (e.key === 'Escape') {{
                window.location.href = 'index.html';
            }}
        }});
    </script>
</body>
</html>"""


def _enhance_complete_html_with_navigation(
    original_html: str,
    slide_number: int,
    total_slides: int,
    topic: str,
    slide_title: str,
) -> str:
    """Enhance complete HTML document with navigation controls"""
    import re

    # Add navigation CSS and JavaScript to the head section
    navigation_css = """
    <style>
        .slide-navigation {
            position: fixed;
            bottom: 20px;
            left: 50%;
            transform: translateX(-50%);
            display: flex;
            gap: 10px;
            z-index: 10000;
            background: rgba(0,0,0,0.8);
            padding: 10px;
            border-radius: 25px;
        }
        .nav-btn {
            background: #007bff;
            color: white;
            border: none;
            padding: 10px 15px;
            border-radius: 5px;
            cursor: pointer;
            text-decoration: none;
            display: inline-block;
            font-size: 14px;
        }
        .nav-btn:hover {
            background: #0056b3;
        }
        .fullscreen-btn {
            position: fixed;
            top: 20px;
            right: 20px;
            background: #28a745;
            color: white;
            border: none;
            padding: 10px;
            border-radius: 5px;
            cursor: pointer;
            z-index: 10000;
            font-size: 16px;
        }
        .fullscreen-btn:hover {
            background: #1e7e34;
        }
    </style>"""

    navigation_js = f"""
    <script>
        function toggleFullscreen() {{
            if (!document.fullscreenElement) {{
                document.documentElement.requestFullscreen();
            }} else {{
                if (document.exitFullscreen) {{
                    document.exitFullscreen();
                }}
            }}
        }}

        // Keyboard navigation
        document.addEventListener('keydown', function(e) {{
            if (e.key === 'ArrowLeft' && {slide_number} > 1) {{
                window.location.href = 'slide_{slide_number-1}.html';
            }} else if (e.key === 'ArrowRight' && {slide_number} < {total_slides}) {{
                window.location.href = 'slide_{slide_number+1}.html';
            }} else if (e.key === 'Escape') {{
                window.location.href = 'index.html';
            }}
        }});
    </script>"""

    navigation_html = f"""
    <div class="slide-navigation">
        <a href="index.html" class="nav-btn">🏠 返回目录</a>
        {"" if slide_number <= 1 else f'<a href="slide_{slide_number-1}.html" class="nav-btn">‹ 上一页</a>'}
        {"" if slide_number >= total_slides else f'<a href="slide_{slide_number+1}.html" class="nav-btn">下一页 ›</a>'}
    </div>

    <button class="fullscreen-btn" onclick="toggleFullscreen()" title="全屏显示">
        📺
    </button>"""

    # Insert navigation CSS into head
    head_pattern = r"</head>"
    enhanced_html = re.sub(
        head_pattern, navigation_css + "\n</head>", original_html, flags=re.IGNORECASE
    )

    # Insert navigation HTML and JS before closing body tag
    body_pattern = r"</body>"
    enhanced_html = re.sub(
        body_pattern,
        navigation_html + "\n" + navigation_js + "\n</body>",
        enhanced_html,
        flags=re.IGNORECASE,
    )

    return enhanced_html


async def _generate_pdf_slide_html(slide, slide_number: int, total_slides: int, topic: str) -> str:
    """Generate PDF-optimized HTML for individual slide without navigation elements"""
    slide_html = slide.get("html_content", "")
    slide_title = slide.get("title", f"第{slide_number}页")

    # Check if it's already a complete HTML document

    if slide_html.strip().lower().startswith("<!doctype") or slide_html.strip().lower().startswith(
        "<html"
    ):
        # It's a complete HTML document, clean it for PDF
        return _clean_html_for_pdf(slide_html, slide_number, total_slides)
    else:
        # It's just content, wrap it in a PDF-optimized structure
        slide_content = slide_html

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{topic} - {slide_title}</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}

        html, body {{
            width: 100%;
            height: 100vh;
            margin: 0;
            padding: 0;
            font-family: 'Microsoft YaHei', 'PingFang SC', sans-serif;
            overflow: hidden;
        }}

        .slide-container {{
            width: 100vw;
            height: 100vh;
            display: flex;
            align-items: center;
            justify-content: center;
            position: relative;
        }}

        .slide-content {{
            width: 100%;
            height: 100%;
            display: flex;
            align-items: center;
            justify-content: center;
            position: relative;
        }}

        /* Ensure all backgrounds and colors are preserved for PDF */
        * {{
            -webkit-print-color-adjust: exact !important;
            print-color-adjust: exact !important;
        }}
    </style>
</head>
<body>
    <div class="slide-container">
        <div class="slide-content">
            {slide_content}
        </div>
    </div>
</body>
</html>"""


def _clean_html_for_pdf(original_html: str, slide_number: int, total_slides: int) -> str:
    """Clean complete HTML document for PDF generation by removing navigation elements"""
    import re

    # Remove navigation elements that might interfere with PDF generation
    cleaned_html = original_html

    # Remove navigation divs and buttons
    cleaned_html = re.sub(
        r'<div[^>]*class="[^"]*navigation[^"]*"[^>]*>.*?</div>',
        "",
        cleaned_html,
        flags=re.DOTALL | re.IGNORECASE,
    )
    cleaned_html = re.sub(
        r'<button[^>]*class="[^"]*nav[^"]*"[^>]*>.*?</button>',
        "",
        cleaned_html,
        flags=re.DOTALL | re.IGNORECASE,
    )
    cleaned_html = re.sub(
        r'<a[^>]*class="[^"]*nav[^"]*"[^>]*>.*?</a>',
        "",
        cleaned_html,
        flags=re.DOTALL | re.IGNORECASE,
    )

    # Remove fullscreen buttons
    cleaned_html = re.sub(
        r"<button[^>]*fullscreen[^>]*>.*?</button>",
        "",
        cleaned_html,
        flags=re.DOTALL | re.IGNORECASE,
    )

    # Replace deprecated placeholder domains with picsum.photos
    cleaned_html = _replace_placeholder_images(cleaned_html)

    # Remove external CDN dependencies for PDF generation (guarded)
    try:
        cleaned_html = _remove_external_dependencies(cleaned_html)
    except Exception:
        # If helper is not available or fails, continue with original html
        pass

    # Add PDF-specific styles
    pdf_styles = """
    <style>
        /* PDF optimization styles */
        * {
            -webkit-print-color-adjust: exact !important;
            print-color-adjust: exact !important;
        }

        html, body {
            width: 100% !important;
            height: 100vh !important;
            margin: 0 !important;
            padding: 0 !important;
            overflow: hidden !important;
        }

        /* Hide any remaining navigation elements */
        .navigation, .nav-btn, .fullscreen-btn, .slide-navigation {
            display: none !important;
        }
    </style>
    """

    # Insert PDF styles before closing head tag
    head_pattern = r"</head>"
    cleaned_html = re.sub(head_pattern, pdf_styles + "\n</head>", cleaned_html, flags=re.IGNORECASE)

    return cleaned_html


def _replace_placeholder_images(html: str) -> str:
    """将常见占位图域名替换为 picsum.photos，尽量保留尺寸"""
    import re

    def repl_img(match):
        before, url, after = match.groups()
        # 尝试提取尺寸
        size_match = re.search(r"(\d{2,4})[xX/](\d{2,4})", url)
        if size_match:
            w, h = size_match.group(1), size_match.group(2)
            new_url = f"https://picsum.photos/{w}/{h}"
        else:
            # 只有一个尺寸或没有尺寸时，默认正方形 400
            one_match = re.search(r"(\d{2,4})(?!\d)", url)
            size = one_match.group(1) if one_match else "400"
            new_url = f"https://picsum.photos/{size}"
        return f"{before}{new_url}{after}"

    # 处理 <img src="...">
    img_pattern = re.compile(
        r'(<img[^>]*\bsrc=")[^"]*(via\.placeholder\.com|placeholder\.com|placehold\.it|dummyimage\.com|placekitten\.com|placekitten|placehold|placeholder)[^"]*("[^>]*>)',
        re.IGNORECASE,
    )
    html = re.sub(img_pattern, repl_img, html)

    # 处理 CSS 背景 url('...')
    def repl_css(match):
        before, url, after = match.groups()
        size_match = re.search(r"(\d{2,4})[xX/](\d{2,4})", url)
        if size_match:
            w, h = size_match.group(1), size_match.group(2)
            new_url = f"https://picsum.photos/{w}/{h}"
        else:
            one_match = re.search(r"(\d{2,4})(?!\d)", url)
            size = one_match.group(1) if one_match else "400"
            new_url = f"https://picsum.photos/{size}"
        return f"{before}{new_url}{after}"

    css_pattern = re.compile(
        r'(url\(["\"])\s*[^)\"\"]*(via\.placeholder\.com|placeholder\.com|placehold\.it|dummyimage\.com|placekitten\.com|placekitten|placehold|placeholder)[^)\"\"]*(\))',
        re.IGNORECASE,
    )
    html = re.sub(css_pattern, repl_css, html)

    return html


def _remove_external_dependencies(html: str) -> str:
    """移除或替换可能影响无头浏览器渲染的外部 CDN 依赖，保证 PDF 生成稳定。
    - 替换 tailwindcdn 脚本为本地样式链接（如果项目提供了本地 CSS）。
    - 移除分析/追踪脚本（analytics/gtag/hotjar 等）。
    - 保留常见功能库，但尽量不破坏页面渲染。
    """
    import re

    cleaned = html

    # 替换 Tailwind CDN 脚本为本地 CSS（若有）
    cleaned = re.sub(
        r'<script[^>]*src=["\']https://cdn\.tailwindcss\.com[^"\']*["\'][^>]*>\s*</script>',
        '<link href="/static/css/tailwind.min.css" rel="stylesheet">',
        cleaned,
        flags=re.IGNORECASE,
    )

    # 移除常见分析/追踪脚本
    patterns = [
        r'<script[^>]*src=["\'][^"\']*analytics[^"\']*["\'][^>]*>\s*</script>',
        r'<script[^>]*src=["\'][^"\']*gtag[^"\']*["\'][^>]*>\s*</script>',
        r'<script[^>]*src=["\'][^"\']*hotjar[^"\']*["\'][^>]*>\s*</script>',
        r'<script[^>]*src=["\'][^"\']*facebook[^"\']*["\'][^>]*>\s*</script>',
        r'<script[^>]*src=["\'][^"\']*twitter[^"\']*["\'][^>]*>\s*</script>',
    ]
    for pat in patterns:
        cleaned = re.sub(pat, '', cleaned, flags=re.IGNORECASE)

    # 替换 jsdelivr CDN 为国内镜像
    cleaned = re.sub(
        r'<script[^>]*src=["\']https://cdn\.jsdelivr\.net/npm/chart\.js[^"\']*["\'][^>]*>\s*</script>',
        '<script src="https://cdn.staticfile.org/chart.js/4.4.0/chart.js"></script>',
        cleaned,
        flags=re.IGNORECASE,
    )

    return cleaned


# duplicate header removed
# ===== Speaker Notes (Script) Generation & Export =====
from ..api.models import SpeakerNotesGenerationRequest


@router.post("/api/projects/{project_id}/scripts/generate")
async def generate_speaker_notes_endpoint(
    project_id: str,
    req: SpeakerNotesGenerationRequest,
    user: User = Depends(get_current_user_required),
):
    """Generate speaker notes for selected or all slides and persist them."""
    project = await ppt_service.project_manager.get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    if not _user_can_access_project(user, project):
        raise HTTPException(status_code=403, detail="Access denied")

    try:
        result = await ppt_service.generate_speaker_notes(project_id, req)
        return JSONResponse({"success": True, **result})
    except Exception as e:
        return JSONResponse(status_code=500, content={"success": False, "error": str(e)})


@router.get("/api/scripts/tasks/{task_id}")
async def get_scripts_task_status(task_id: str, user: User = Depends(get_current_user_required)):
    task = ppt_service.scripts_tasks.get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    return JSONResponse({"success": True, **task, "task_id": task_id})

@router.put("/api/projects/{project_id}/slides/{slide_number}/speaker_notes")
async def update_single_speaker_notes(project_id: str, slide_number: int, payload: dict, user: User = Depends(get_current_user_required)):
    """Update speaker notes for one slide (slide_number 1-based)."""
    project = await ppt_service.project_manager.get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    if not _user_can_access_project(user, project):
        raise HTTPException(status_code=403, detail="Access denied")
    idx = slide_number - 1
    if idx < 0 or idx >= len(project.slides_data or []):
        raise HTTPException(status_code=400, detail="Invalid slide number")
    notes = (payload.get('notes') or '').strip()
    slide = project.slides_data[idx] or {}
    meta = slide.get('metadata') or {}
    meta['speaker_notes'] = notes
    slide['metadata'] = meta
    slide['speaker_notes'] = notes
    project.slides_data[idx] = slide
    await ppt_service.project_manager.save_single_slide(project_id, idx, slide)
    await ppt_service.project_manager.update_project_data(project_id, {"slides_data": project.slides_data})
    return JSONResponse({"success": True, "index": slide_number, "length": len(notes)})


def _build_notes_markdown(project, indices: list[int] | None = None) -> str:
    """Assemble Markdown with titles and notes for selected slides."""
    lines = []
    lines.append(f"# {project.topic} - 演讲稿")
    slides = project.slides_data or []
    if not slides:
        return "# 演讲稿\n\n（暂无内容）"

    chosen = range(len(slides)) if not indices else [i for i in indices if 0 <= i < len(slides)]
    for i in chosen:
        s = slides[i]
        title = s.get("title", f"第{i+1}页")
        notes = (s.get("metadata", {}).get("speaker_notes") or s.get("speaker_notes") or "").strip()
        lines.append("")
        lines.append(f"## 第{i+1}页 · {title}")
        if notes:
            lines.append("")
            lines.append(notes)
        else:
            lines.append("")
            lines.append("（暂无演讲稿）")
    return "\n".join(lines)


@router.get("/api/projects/{project_id}/export/scripts.md")
async def export_speaker_notes_markdown(
    project_id: str,
    indices: Optional[str] = None,  # comma-separated 1-based
    include_meta: Optional[str] = "0",  # accept flexible truthy values
    user: User = Depends(get_current_user_required),
):
    """Export speaker notes as a Markdown file for selected or all slides."""
    project = await ppt_service.project_manager.get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    if not _user_can_access_project(user, project):
        raise HTTPException(status_code=403, detail="Access denied")

    idx0 = None
    if indices:
        try:
            idx0 = [int(x.strip()) - 1 for x in indices.split(",") if x.strip()]
        except Exception:
            raise HTTPException(status_code=400, detail="Invalid indices parameter")

    md = _build_notes_markdown(project, idx0)

    # Normalize include_meta flag (supports 1/true/yes/on)
    flag = False
    try:
        if isinstance(include_meta, int):
            flag = include_meta == 1
        elif isinstance(include_meta, str):
            flag = include_meta.strip().lower() in {"1", "true", "yes", "on"}
    except Exception:
        flag = False

    # Optional metadata header (YAML-like front matter)
    if flag:
        try:
            total_slides = len(project.slides_data or [])
            # 统计已有演讲稿页数
            notes_count = 0
            for s in (project.slides_data or []):
                note_txt = (s.get("metadata", {}).get("speaker_notes") or s.get("speaker_notes") or "").strip()
                if note_txt:
                    notes_count += 1
            # 可扩展的统计：如果后续任务结构里存储了生成策略，可在这里补充
            meta_lines = [
                "---",
                f"topic: {project.topic}",
                f"slides_total: {total_slides}",
                f"slides_with_notes: {notes_count}",
                f"export_time: {time.strftime('%Y-%m-%d %H:%M:%S', time.localtime())}",
                "generator: flowslide",
                "metadata_included: true",
                "format_version: 1",
                "---",
                "",
            ]
            md = "\n".join(meta_lines) + md
        except Exception:
            pass

    # Prepare response (avoid non-ASCII in plain filename= to prevent latin-1 header encoding errors)
    # 1) Build a human-readable Unicode filename for filename*
    unicode_topic = re.sub(r"[^\w\s-]", "", project.topic).strip() or "presentation"
    unicode_filename = f"{unicode_topic}_演讲稿.md"
    # 2) Build an ASCII-safe fallback for filename=
    ascii_topic = re.sub(r"[^A-Za-z0-9_-]", "", unicode_topic)
    if not ascii_topic:
        ascii_topic = "presentation"
    ascii_filename = f"{ascii_topic}_notes.md"
    encoded = urllib.parse.quote(unicode_filename, safe="")
    from fastapi.responses import Response

    return Response(
        content=md,
        media_type="text/markdown; charset=utf-8",
        headers={
            # Provide ASCII fallback and RFC 5987 UTF-8 filename*
            "Content-Disposition": f"attachment; filename=\"{ascii_filename}\"; filename*=UTF-8''{encoded}"
        },
    )


@router.get("/api/projects/{project_id}/export/scripts.docx")
async def export_speaker_notes_docx(
    project_id: str,
    indices: Optional[str] = None,  # comma-separated 1-based
    include_meta: Optional[str] = "0",
    user: User = Depends(get_current_user_required),
):
    """Export speaker notes as a DOCX file using python-docx."""
    project = await ppt_service.project_manager.get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    if not _user_can_access_project(user, project):
        raise HTTPException(status_code=403, detail="Access denied")

    idx0 = None
    if indices:
        try:
            idx0 = [int(x.strip()) - 1 for x in indices.split(",") if x.strip()]
        except Exception:
            raise HTTPException(status_code=400, detail="Invalid indices parameter")

    # Build DOCX in a temp file
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception:
        raise HTTPException(status_code=503, detail="python-docx is not installed")

    slides = project.slides_data or []
    chosen = range(len(slides)) if not idx0 else [i for i in idx0 if 0 <= i < len(slides)]

    document = Document()
    document.core_properties.title = f"{project.topic} - 演讲稿"

    # Title + optional metadata
    heading = document.add_heading(f"{project.topic} - 演讲稿", level=0)
    # Parse include_meta for docx same as markdown
    flag_docx = False
    try:
        if isinstance(include_meta, int):
            flag_docx = include_meta == 1
        elif isinstance(include_meta, str):
            flag_docx = include_meta.strip().lower() in {"1", "true", "yes", "on"}
    except Exception:
        flag_docx = False
    if flag_docx:
        try:
            total_slides = len(slides)
            notes_count = 0
            for s in slides:
                txt = (s.get("metadata", {}).get("speaker_notes") or s.get("speaker_notes") or "").strip()
                if txt:
                    notes_count += 1
            meta_para = document.add_paragraph()
            meta_para.add_run(f"导出时间: {time.strftime('%Y-%m-%d %H:%M:%S', time.localtime())}\n")
            meta_para.add_run(f"总页数: {total_slides}  有演讲稿页数: {notes_count}\n")
            meta_para.add_run("生成器: flowslide  (metadata_included=true, format_version=1)\n")
        except Exception:
            pass

    for i in chosen:
        s = slides[i]
        title = s.get("title", f"第{i+1}页")
        notes = (s.get("metadata", {}).get("speaker_notes") or s.get("speaker_notes") or "").strip()

        document.add_heading(f"第{i+1}页 · {title}", level=1)
        if notes:
            # Split into paragraphs on blank lines; fall back to single paragraph
            parts = [p.strip() for p in re.split(r"\n\s*\n", notes) if p.strip()]
            if not parts:
                parts = [notes]
            for p in parts:
                para = document.add_paragraph(p)
                para_format = para.paragraph_format
                para_format.space_after = Pt(6)
        else:
            document.add_paragraph("（暂无演讲稿）")

    # Save to temp path
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        temp_path = tmp.name
    document.save(temp_path)

    unicode_topic = re.sub(r"[^\w\s-]", "", project.topic).strip()[:50] or "presentation"
    unicode_filename = f"{unicode_topic}_演讲稿.docx"
    ascii_topic = re.sub(r"[^A-Za-z0-9_-]", "", unicode_topic)
    if not ascii_topic:
        ascii_topic = "presentation"
    ascii_filename = f"{ascii_topic}_notes.docx"  # ASCII-safe fallback
    encoded = urllib.parse.quote(unicode_filename, safe="")
    from starlette.background import BackgroundTask

    def cleanup():
        try:
            os.unlink(temp_path)
        except Exception:
            pass

    return FileResponse(
        temp_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        # Don't set filename param to avoid duplicate Content-Disposition; provide our own header with ASCII + filename*
        headers={
            "Content-Disposition": f"attachment; filename=\"{ascii_filename}\"; filename*=UTF-8''{encoded}"
        },
        background=BackgroundTask(cleanup),
    )


async def _generate_pdf_with_pyppeteer(project, output_path: str, individual: bool = False) -> bool:
    """Generate PDF using Pyppeteer (Python)"""
    try:
        pdf_converter = get_pdf_converter()

        # Check if PDF converter is available
        if not pdf_converter.is_available():
            logging.error("PDF converter is not available - Playwright not installed")
            return False

        logging.info(f"Starting PDF generation for project: {project.topic}")
        logging.info(f"Number of slides: {len(project.slides_data)}")

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            logging.info(f"Using temporary directory: {temp_path}")

            # Always generate individual HTML files for each slide for better page separation
            # This ensures each slide becomes a separate PDF page
            html_files = []
            for i, slide in enumerate(project.slides_data):
                try:
                    logging.info(f"Processing slide {i+1}: {slide.get('title', 'No title')}")

                    # Check slide data structure
                    if not isinstance(slide, dict):
                        logging.error(f"Slide {i+1} is not a dictionary: {type(slide)}")
                        return False

                    if "html_content" not in slide:
                        logging.error(
                            f"Slide {i+1} missing html_content key. Keys: {list(slide.keys())}"
                        )
                        return False

                    html_content = slide.get("html_content", "")
                    if not html_content or not html_content.strip():
                        logging.error(f"Slide {i+1} has empty html_content")
                        return False

                    logging.info(f"Slide {i+1} html_content length: {len(html_content)} characters")

                    # Use a specialized PDF-optimized HTML generator without navigation
                    slide_html = await _generate_pdf_slide_html(
                        slide, i + 1, len(project.slides_data), project.topic
                    )

                    if not slide_html or not slide_html.strip():
                        logging.error(f"Generated HTML for slide {i+1} is empty")
                        return False

                    html_file = temp_path / f"slide_{i+1}.html"

                    # Write HTML file in thread pool to avoid blocking
                    def write_html_file(content, path):
                        with open(path, "w", encoding="utf-8") as f:
                            f.write(content)

                    await run_blocking_io(write_html_file, slide_html, str(html_file))

                    # Verify file was written
                    if not html_file.exists():
                        logging.error(f"HTML file was not created: {html_file}")
                        return False

                    file_size = html_file.stat().st_size
                    if file_size == 0:
                        logging.error(f"HTML file is empty: {html_file}")
                        return False

                    html_files.append(str(html_file))
                    logging.info(
                        f"Generated HTML file for slide {i+1}: {html_file} ({file_size} bytes)"
                    )

                except Exception as e:
                    logging.error(f"Failed to generate HTML for slide {i+1}: {e}", exc_info=True)
                    return False

            if not html_files:
                logging.error("No HTML files were generated")
                return False

            # Use Pyppeteer to convert multiple files and merge them
            pdf_dir = temp_path / "pdfs"
            await run_blocking_io(pdf_dir.mkdir)

            logging.info(f"Starting PDF generation for {len(html_files)} files")

            # Convert HTML files to PDFs and merge them
            pdf_files = await pdf_converter.convert_multiple_html_to_pdf(
                html_files, str(pdf_dir), output_path
            )

            if pdf_files and os.path.exists(output_path):
                file_size = os.path.getsize(output_path)
                logging.info(f"PDF generation successful: {output_path} ({file_size} bytes)")
                return True
            else:
                logging.error("PDF generation failed: No output file created or file is empty")
                if os.path.exists(output_path):
                    file_size = os.path.getsize(output_path)
                    logging.error(f"Output file exists but size is: {file_size} bytes")
                return False

    except Exception as e:
        logging.error(f"PDF generation failed with exception: {e}", exc_info=True)
        return False


async def _generate_combined_html_for_pdf(project) -> str:
    """Generate combined HTML for PDF export with all slides preserving original styles"""
    slides_html = ""
    global_styles = ""

    for i, slide in enumerate(project.slides_data):
        slide_html = slide.get("html_content", "")
        slide_title = slide.get("title", f"第{i+1}页")

        # Enhanced style extraction to preserve all styling
        if slide_html.strip().lower().startswith(
            "<!doctype"
        ) or slide_html.strip().lower().startswith("<html"):
            import re

            # Extract all CSS styles from head (including link tags and style tags)
            style_matches = re.findall(
                r"<style[^>]*>(.*?)</style>", slide_html, re.DOTALL | re.IGNORECASE
            )
            link_matches = re.findall(
                r'<link[^>]*rel=["\']stylesheet["\'][^>]*>', slide_html, re.IGNORECASE
            )

            slide_styles = "\n".join(style_matches)
            slide_links = "\n".join(link_matches)

            # Extract body content with preserved attributes
            body_match = re.search(
                r"<body([^>]*)>(.*?)</body>", slide_html, re.DOTALL | re.IGNORECASE
            )
            if body_match:
                body_attrs = body_match.group(1)
                slide_content = body_match.group(2)
                # Preserve body styles if any
                if "style=" in body_attrs:
                    body_style_match = re.search(r'style=["\']([^"\']*)["\']', body_attrs)
                    if body_style_match:
                        slide_styles += f"\n.slide-content {{ {body_style_match.group(1)} }}"
            else:
                slide_content = slide_html
                slide_links = ""

            # Add to global styles to avoid duplication
            if slide_links and slide_links not in global_styles:
                global_styles += slide_links + "\n"
        else:
            slide_styles = ""
            slide_content = slide_html
            slide_links = ""

        # Create a separate page for each slide with proper page break
        slides_html += f"""
        <div class="slide-page" data-slide="{i+1}" style="page-break-before: always; page-break-after: always; page-break-inside: avoid;">
            <style>
                /* Slide {i+1} specific styles */
                .slide-page[data-slide="{i+1}"] .slide-content {{
                    /* Preserve original styling */
                }}
                {slide_styles}
            </style>
            <div class="slide-content">
                {slide_content}
            </div>
            <div class="slide-footer">
                <span class="slide-number">{i+1} / {len(project.slides_data)}</span>
                <span class="slide-title">{slide_title}</span>
            </div>
        </div>"""

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{project.topic} - PDF导出</title>
    {global_styles}
    <style>
        /* Reset and base styles */
        * {{
            box-sizing: border-box;
        }}

        body {{
            margin: 0;
            padding: 0;
            font-family: 'Microsoft YaHei', 'PingFang SC', sans-serif;
            /* Don't force background color - let slides define their own */
        }}

        .slide-page {{
            width: 297mm;
            height: 167mm;
            margin: 0;
            padding: 0;
            page-break-before: always;
            page-break-after: always;
            page-break-inside: avoid;
            position: relative;
            aspect-ratio: 16/9;
            /* Don't force background - preserve original slide backgrounds */
            overflow: hidden;
            display: block;
            box-sizing: border-box;
        }}

        .slide-page:first-child {{
            page-break-before: avoid;
        }}

        .slide-page:last-child {{
            page-break-after: avoid;
        }}

        .slide-content {{
            width: 100%;
            height: calc(100% - 30px);
            position: relative;
            /* Preserve original content styling */
        }}

        .slide-footer {{
            position: absolute;
            bottom: 5mm;
            right: 10mm;
            font-size: 10px;
            color: rgba(255, 255, 255, 0.7);
            background: rgba(0, 0, 0, 0.3);
            padding: 2px 8px;
            border-radius: 3px;
            z-index: 1000;
        }}

        .slide-number {{
            font-weight: bold;
        }}

        .slide-title {{
            margin-left: 8px;
            opacity: 0.8;
        }}

        /* Print-specific styles */
        @media print {{
            @page {{
                size: 297mm 167mm;
                margin: 0;
            }}

            body {{
                -webkit-print-color-adjust: exact;
                print-color-adjust: exact;
                margin: 0;
                padding: 0;
            }}

            .slide-page {{
                page-break-before: always;
                page-break-after: always;
                page-break-inside: avoid;
                -webkit-print-color-adjust: exact;
                print-color-adjust: exact;
                width: 297mm;
                height: 167mm;
                margin: 0;
                padding: 0;
                display: block;
            }}

            .slide-page:first-child {{
                page-break-before: avoid;
            }}

            .slide-page:last-child {{
                page-break-after: avoid;
            }}
        }}

        /* Ensure all backgrounds and colors are preserved */
        * {{
            -webkit-print-color-adjust: exact !important;
            print-color-adjust: exact !important;
        }}
    </style>
</head>
<body>
    {slides_html}
</body>
</html>"""


async def _generate_slideshow_index(project, slide_files: list) -> str:
    """Generate slideshow index page"""
    slides_list = ""
    for i, slide_file in enumerate(slide_files):
        slide = project.slides_data[i]
        slide_title = slide.get("title", f"第{i+1}页")
        slides_list += f"""
        <div class="slide-item" onclick="openSlide('{slide_file}')">
            <div class="slide-preview">
                <div class="slide-number">{i+1}</div>
                <div class="slide-title">{slide_title}</div>
            </div>
        </div>"""

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{project.topic} - Slide放映</title>
    <style>
        body {{
            margin: 0;
            padding: 0;
            font-family: 'Microsoft YaHei', 'PingFang SC', sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
        }}
        .header {{
            text-align: center;
            padding: 40px 20px;
            color: white;
        }}
        .header h1 {{
            margin: 0;
            font-size: 2.5em;
            font-weight: 300;
        }}
        .header p {{
            margin: 10px 0 0 0;
            font-size: 1.2em;
            opacity: 0.9;
        }}
        .container {{
            max-width: 1200px;
            margin: 0 auto;
            padding: 0 20px;
        }}
        .slides-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fill, minmax(300px, 1fr));
            gap: 20px;
            padding: 20px 0;
        }}
        .slide-item {{
            background: white;
            border-radius: 10px;
            padding: 20px;
            cursor: pointer;
            transition: all 0.3s ease;
            box-shadow: 0 4px 15px rgba(0,0,0,0.1);
        }}
        .slide-item:hover {{
            transform: translateY(-5px);
            box-shadow: 0 8px 25px rgba(0,0,0,0.2);
        }}
        .slide-preview {{
            text-align: center;
        }}
        .slide-number {{
            background: #007bff;
            color: white;
            width: 40px;
            height: 40px;
            border-radius: 50%;
            display: flex;
            align-items: center;
            justify-content: center;
            margin: 0 auto 15px auto;
            font-weight: bold;
        }}
        .slide-title {{
            font-size: 1.1em;
            color: #333;
            margin: 0;
        }}
        .controls {{
            text-align: center;
            padding: 40px 20px;
        }}
        .btn {{
            background: #28a745;
            color: white;
            border: none;
            padding: 15px 30px;
            border-radius: 25px;
            font-size: 1.1em;
            cursor: pointer;
            margin: 0 10px;
            transition: all 0.3s ease;
            text-decoration: none;
            display: inline-block;
        }}
        .btn:hover {{
            background: #1e7e34;
            transform: translateY(-2px);
        }}
        .btn-secondary {{
            background: #6c757d;
        }}
        .btn-secondary:hover {{
            background: #545b62;
        }}
        @media (max-width: 768px) {{
            .slides-grid {{
                grid-template-columns: repeat(auto-fill, minmax(250px, 1fr));
                gap: 15px;
            }}
            .header h1 {{
                font-size: 2em;
            }}
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{project.topic}</h1>
        <p>Slide演示文稿 - 共{len(slide_files)}页</p>
    </div>

    <div class="container">
        <div class="controls">
            <button class="btn" onclick="startSlideshow()">🎬 开始放映</button>
            <button class="btn btn-secondary" onclick="downloadAll()">📦 下载所有文件</button>
        </div>

        <div class="slides-grid">
            {slides_list}
        </div>
    </div>

    <script>
        function openSlide(slideFile) {{
            window.open(slideFile, '_blank');
        }}

        function startSlideshow() {{
            window.open('slide_1.html', '_blank');
        }}

        function downloadAll() {{
            alert('所有文件已包含在此ZIP包中');
        }}

        // Keyboard shortcuts
        document.addEventListener('keydown', function(e) {{
            if (e.key === 'Enter' || e.key === ' ') {{
                startSlideshow();
            }}
        }});
    </script>
</body>
</html>"""


@router.get("/upload", response_class=HTMLResponse)
async def web_upload_page(request: Request, user: User = Depends(get_current_user_required)):
    """File upload page"""
    return templates.TemplateResponse("upload.html", {"request": request})


async def _process_uploaded_file_for_outline(
    file_upload: UploadFile,
    topic: str,
    target_audience: str,
    page_count_mode: str,
    min_pages: int,
    max_pages: int,
    fixed_pages: int,
    ppt_style: str,
    custom_style_prompt: str,
    file_processing_mode: str,
    content_analysis_depth: str,
    requirements: str = None,
) -> Optional[Dict[str, Any]]:
    """处理上传的文件并生成Slide大纲"""
    try:
        # 验证文件
        from ..services.file_processor import FileProcessor

        file_processor = FileProcessor()

        is_valid, message = file_processor.validate_file(file_upload.filename, file_upload.size)
        if not is_valid:
            logger.error(f"File validation failed: {message}")
            return None

        # 读取文件内容并保存临时文件（在线程池中执行）
        content = await file_upload.read()
        temp_file_path = await run_blocking_io(_save_temp_file_sync, content, file_upload.filename)

        try:
            # 创建文件大纲生成请求
            from ..api.models import FileOutlineGenerationRequest

            outline_request = FileOutlineGenerationRequest(
                file_path=temp_file_path,
                filename=file_upload.filename,
                topic=topic if topic.strip() else None,
                scenario="general",  # 默认场景，可以根据需要调整
                requirements=requirements,
                target_audience=target_audience,
                page_count_mode=page_count_mode,
                min_pages=min_pages,
                max_pages=max_pages,
                fixed_pages=fixed_pages,
                ppt_style=ppt_style,
                custom_style_prompt=custom_style_prompt,
                file_processing_mode=file_processing_mode,
                content_analysis_depth=content_analysis_depth,
            )

            # 使用enhanced_ppt_service生成大纲
            result = await ppt_service.generate_outline_from_file(outline_request)

            if result.success:
                logger.info(f"Successfully generated outline from file: {file_upload.filename}")
                return result.outline
            else:
                logger.error(f"Failed to generate outline from file: {result.error}")
                return None

        finally:
            # 清理临时文件（在线程池中执行）
            await run_blocking_io(_cleanup_temp_file_sync, temp_file_path)

    except Exception as e:
        logger.error(f"Error processing uploaded file for outline: {e}")
        return None


def _save_temp_file_sync(content: bytes, filename: str) -> str:
    """同步保存临时文件（在线程池中运行）"""
    import os
    import tempfile

    with tempfile.NamedTemporaryFile(
        delete=False, suffix=os.path.splitext(filename)[1]
    ) as temp_file:
        temp_file.write(content)
        return temp_file.name


def _cleanup_temp_file_sync(temp_file_path: str):
    """同步清理临时文件（在线程池中运行）"""
    import os

    if os.path.exists(temp_file_path):
        os.unlink(temp_file_path)


@router.get("/global-master-templates", response_class=HTMLResponse)
async def global_master_templates_page(
    request: Request, user: User = Depends(get_current_user_required)
):
    """Global master templates management page"""
    try:
        return templates.TemplateResponse(
            "global_master_templates.html",
            {"request": request, "aspect_ratio_settings": get_aspect_ratio_settings()},
        )
    except Exception as e:
        logger.error(f"Error loading global master templates page: {e}")
        return templates.TemplateResponse("error.html", {"request": request, "error": str(e)})


@router.get("/image-gallery", response_class=HTMLResponse)
async def image_gallery_page(request: Request, user: User = Depends(get_current_user_required)):
    """本地图床管理页面"""
    try:
        return templates.TemplateResponse("image_gallery.html", {"request": request, "user": user})
    except Exception as e:
        logger.error(f"Error rendering image gallery page: {e}")
        return templates.TemplateResponse("error.html", {"request": request, "error": str(e)})


@router.get("/image-generation-test", response_class=HTMLResponse)
async def image_generation_test_page(
    request: Request, user: User = Depends(get_current_user_required)
):
    """AI图片生成测试页面"""
    try:
        return templates.TemplateResponse(
            "image_generation_test.html", {"request": request, "user": user}
        )
    except Exception as e:
        logger.error(f"Error rendering image generation test page: {e}")
        return templates.TemplateResponse("error.html", {"request": request, "error": str(e)})


@router.get("/projects/{project_id}/template-selection", response_class=HTMLResponse)
async def template_selection_page(
    request: Request, project_id: str, user: User = Depends(get_current_user_required)
):
    """Template selection page for Slide generation"""
    try:
        # Get project info
        project = await ppt_service.project_manager.get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")

        return templates.TemplateResponse(
            "template_selection.html",
            {
                "request": request,
                "project_id": project_id,
                "project_topic": project.topic,
                "aspect_ratio_settings": get_aspect_ratio_settings(),
            },
        )
    except Exception as e:
        logger.error(f"Error loading template selection page: {e}")
        return templates.TemplateResponse("error.html", {"request": request, "error": str(e)})


# 图像重新生成相关辅助函数
async def analyze_image_context(
    image_info: Dict[str, Any],
    slide_content: Dict[str, Any],
    project_topic: str,
    project_scenario: str,
) -> Dict[str, Any]:
    """分析图像在幻灯片中的上下文"""
    return {
        "slide_title": slide_content.get("title", ""),
        "slide_content": slide_content.get("html_content", ""),
        "image_alt": image_info.get("alt", ""),
        "image_title": image_info.get("title", ""),
        "image_size": f"{image_info.get('width', 0)}x{image_info.get('height', 0)}",
        "image_position": image_info.get("position", {}),
        "project_topic": project_topic,
        "project_scenario": project_scenario,
        "image_purpose": determine_image_purpose(image_info, slide_content),
    }


def determine_image_purpose(image_info: Dict[str, Any], slide_content: Dict[str, Any]) -> str:
    """确定图像在幻灯片中的用途"""
    # 简单的启发式规则来确定图像用途
    width = image_info.get("width", 0)
    height = image_info.get("height", 0)
    alt_text = image_info.get("alt", "").lower()

    if width > 800 or height > 600:
        return "background"  # 大图像可能是背景
    elif "icon" in alt_text or "logo" in alt_text:
        return "icon"
    elif "chart" in alt_text or "graph" in alt_text:
        return "chart_support"
    elif width < 200 and height < 200:
        return "decoration"
    else:
        return "illustration"


# 图像重新生成相关辅助函数


def select_best_image_source(
    enabled_sources: List, image_config: Dict[str, Any], image_context: Dict[str, Any]
):
    """智能选择最佳的图片来源"""
    from ..services.models.slide_image_info import ImageSource

    # 如果只有一个启用的来源，直接使用
    if len(enabled_sources) == 1:
        return enabled_sources[0]

    # 根据图像用途和配置智能选择
    image_purpose = image_context.get("image_purpose", "illustration")

    # 优先级规则
    if image_purpose == "background":
        # 背景图优先使用AI生成，其次网络搜索
        if ImageSource.AI_GENERATED in enabled_sources:
            return ImageSource.AI_GENERATED
        elif ImageSource.NETWORK in enabled_sources:
            return ImageSource.NETWORK
        elif ImageSource.LOCAL in enabled_sources:
            return ImageSource.LOCAL

    elif image_purpose == "icon":
        # 图标优先使用本地，其次AI生成
        if ImageSource.LOCAL in enabled_sources:
            return ImageSource.LOCAL
        elif ImageSource.AI_GENERATED in enabled_sources:
            return ImageSource.AI_GENERATED
        elif ImageSource.NETWORK in enabled_sources:
            return ImageSource.NETWORK

    elif image_purpose in ["illustration", "chart_support", "decoration"]:
        # 说明性图片优先使用网络搜索，其次AI生成
        if ImageSource.NETWORK in enabled_sources:
            return ImageSource.NETWORK
        elif ImageSource.AI_GENERATED in enabled_sources:
            return ImageSource.AI_GENERATED
        elif ImageSource.LOCAL in enabled_sources:
            return ImageSource.LOCAL

    # 默认优先级：AI生成 > 网络搜索 > 本地
    for source in [ImageSource.AI_GENERATED, ImageSource.NETWORK, ImageSource.LOCAL]:
        if source in enabled_sources:
            return source

    # 如果都没有，返回第一个可用的
    return enabled_sources[0] if enabled_sources else ImageSource.AI_GENERATED


# 注意：generate_image_prompt_for_replacement 函数已被PPTImageProcessor的标准流程替代
# 现在使用 PPTImageProcessor._ai_generate_image_prompt 方法来生成提示词


def replace_image_in_html(html_content: str, image_info: Dict[str, Any], new_image_url: str) -> str:
    """在HTML内容中替换指定的图像，支持img标签、背景图像和SVG，保持布局和样式"""
    try:

        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html_content, "html.parser")

        old_src = image_info.get("src", "")
        image_type = image_info.get("type", "img")

        if not old_src:
            logger.warning("图像信息中没有src属性，无法替换")
            return html_content

        replacement_success = False

        if image_type == "img":
            # 处理 <img> 标签
            replacement_success = replace_img_tag(soup, image_info, new_image_url, old_src)

        elif image_type == "background":
            # 处理背景图像
            replacement_success = replace_background_image(soup, image_info, new_image_url, old_src)

        elif image_type == "svg":
            # 处理SVG图像
            replacement_success = replace_svg_image(soup, image_info, new_image_url, old_src)

        if replacement_success:
            logger.info(f"成功替换{image_type}图像: {old_src} -> {new_image_url}")
            return str(soup)
        else:
            logger.warning(f"未找到匹配的{image_type}图像进行替换")
            return fallback_string_replacement(html_content, old_src, new_image_url)

    except Exception as e:
        logger.error(f"替换HTML中的图像失败: {e}")
        return fallback_string_replacement(html_content, image_info.get("src", ""), new_image_url)


def replace_img_tag(soup, image_info: Dict[str, Any], new_image_url: str, old_src: str) -> bool:
    """替换img标签"""
    img_elements = soup.find_all("img")

    for img in img_elements:
        img_src = img.get("src", "")

        # 比较图像源URL（处理相对路径和绝对路径）
        if (
            img_src == old_src
            or img_src.endswith(old_src.split("/")[-1])
            or old_src.endswith(img_src.split("/")[-1])
        ):

            # 替换图像URL
            img["src"] = new_image_url

            # 保持原有的重要属性
            preserved_attributes = ["class", "style", "width", "height", "id"]
            for attr in preserved_attributes:
                if attr in image_info and image_info[attr]:
                    img[attr] = image_info[attr]

            # 更新或保持alt和title
            if image_info.get("alt"):
                img["alt"] = image_info["alt"]
            if image_info.get("title"):
                img["title"] = image_info["title"]

            # 确保图像加载错误时有后备处理
            if not img.get("onerror"):
                img["onerror"] = "this.style.display='none'"

            return True

    return False


def replace_background_image(
    soup, image_info: Dict[str, Any], new_image_url: str, old_src: str
) -> bool:
    """替换CSS背景图像"""
    # 查找所有元素
    all_elements = soup.find_all()

    for element in all_elements:
        # 检查内联样式中的背景图像
        style = element.get("style", "")
        if "background-image" in style and old_src in style:
            # 替换内联样式中的背景图像URL
            new_style = style.replace(old_src, new_image_url)
            element["style"] = new_style
            return True

        # 检查class属性，可能对应CSS规则中的背景图像
        class_names = element.get("class", [])
        if class_names and image_info.get("className"):
            # 如果class匹配，我们假设这是目标元素
            if any(cls in image_info.get("className", "") for cls in class_names):
                # 为元素添加内联背景图像样式
                current_style = element.get("style", "")
                if current_style and not current_style.endswith(";"):
                    current_style += ";"
                new_style = f"{current_style}background-image: url('{new_image_url}');"
                element["style"] = new_style
                return True

    return False


def replace_svg_image(soup, image_info: Dict[str, Any], new_image_url: str, old_src: str) -> bool:
    """替换SVG图像"""
    # 查找SVG元素
    svg_elements = soup.find_all("svg")

    for svg in svg_elements:
        # 如果SVG有src属性（虽然不常见）
        if svg.get("src") == old_src:
            svg["src"] = new_image_url
            return True

        # 检查SVG的内容或其他标识
        if image_info.get("outerHTML") and svg.get_text() in image_info.get("outerHTML", ""):
            # 对于内联SVG，我们可能需要替换整个元素
            # 这里简化处理，添加一个data属性来标记已替换
            svg["data-replaced-image"] = new_image_url
            return True

    return False


def fallback_string_replacement(html_content: str, old_src: str, new_image_url: str) -> str:
    """后备的字符串替换方案"""
    try:
        import re

        if old_src and old_src in html_content:
            # 尝试多种替换模式
            patterns = [
                # img标签的src属性
                (
                    rf'(<img[^>]*src=")[^"]*({re.escape(old_src)}[^"]*")([^>]*>)',
                    rf"\1{new_image_url}\3",
                ),
                # CSS背景图像
                (
                    rf'(background-image:\s*url\([\'"]?)[^\'")]*({re.escape(old_src)}[^\'")]*)',
                    rf"\1{new_image_url}",
                ),
                # 直接字符串替换
                (re.escape(old_src), new_image_url),
            ]

            for pattern, replacement in patterns:
                updated_html = re.sub(pattern, replacement, html_content, flags=re.IGNORECASE)
                if updated_html != html_content:
                    logger.info(f"使用后备方案成功替换图像: {old_src} -> {new_image_url}")
                    return updated_html

        return html_content

    except Exception as e:
        logger.error(f"后备替换方案也失败: {e}")
        return html_content
